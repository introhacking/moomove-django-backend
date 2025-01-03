# views.py in your Django app
from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import status
import pandas as pd
import numpy as np
from datetime import datetime
from .models import *
from .serializers import *
from pdfplumber import open as open_pdf
import PyPDF2
from django.http import JsonResponse
from django.views import View
from django.views.decorators.csrf import csrf_exempt
import re
# import fitz
from django.db import transaction
from rest_framework.parsers import MultiPartParser, FormParser, JSONParser
from docx import Document
from rest_framework import generics
from django.shortcuts import get_object_or_404
from django.contrib.auth import authenticate
from rest_framework_simplejwt.tokens import RefreshToken
from rest_framework.decorators import api_view
from rest_framework.permissions import IsAuthenticated
from drf_yasg import openapi

from rest_framework.permissions import AllowAny
from drf_spectacular.utils import extend_schema
from drf_yasg.utils import swagger_auto_schema
from google.oauth2 import service_account
from google.cloud import documentai_v1beta3 as documentai
from django.db import connection
from decimal import Decimal
import uuid
from dotenv import load_dotenv, dotenv_values
from rest_framework.authtoken.models import Token

# from .permissions import IsSystemAdmin, IsClientAdmin, IsClientUser
from rest_framework.exceptions import ValidationError
from uauth.serializers import *
from uauth.models import *
from uauth.views import *
from uauth.role_permission import *
from datetime import datetime, timedelta
from django.db.models import Q

config={
    **dotenv_values('constant_env/.env.shared'),
    **dotenv_values('constant_env/.env.secret'),
    **dotenv_values('constant_env/.env.error'),
}

# required_columns = [item.strip() for item in config.get("REQUIRED_COLUMN", '').split(",")]
# print(required_columns)
# print(config.get("REQUIRED_COLUMN" , "").split(","))
@extend_schema(
    request=UserSerializer,
    responses={200: 'application/json'},
    tags=["Authentication"],
    summary="Login to get JWT token",
    description="Provide username and password to receive JWT access and refresh tokens."
)


# ROLE VIEWS
# class RegisterView(APIView):
#     permission_classes = [AllowAny]

#     def post(self, request):
#         serializer = RegisterSerializer(data=request.data)
#         if serializer.is_valid():
#             user = serializer.save()

#             # Generate a token for the user (optional if you're using Token Authentication)
#             from rest_framework.authtoken.models import Token
#             token = Token.objects.create(user=user)

#             return Response({
#                 "message": "User registered successfully",
#                 "token": token.key
#             })
        
#         return Response(serializer.errors, status=400)


# class AdminDashboardView(APIView):
#     permission_classes = [IsSystemAdmin]

#     def get(self, request):
#         return Response({"message": config['SYSTEM_MESSAGE']})

# class ClientDashboardView(APIView):
#     permission_classes = [IsClientAdmin]

#     def get(self, request):
#         return Response({"message": config['CLIENT_MESSAGE']})

# class GenerateQuoteView(APIView):
#     permission_classes = [IsClientUser]

#     def post(self, request):
#         return Response({"message": "Quote generated successfully!"})



# @api_view(['POST' , 'GET'])
# def login(request):
#     serializer = UserSerializer(data=request.data)
#     if serializer.is_valid():
#         username = serializer.validated_data['username']
#         password = serializer.validated_data['password']
#         user = authenticate(request, username=username, password=password)
#         if user is not None:
#             refresh = RefreshToken.for_user(user)
#             return Response({
#                 'refresh': str(refresh),
#                 'access': str(refresh.access_token),
#             })
#         else:
#             return Response({'detail': 'Invalid credentials'}, status=status.HTTP_401_UNAUTHORIZED)
#     return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

# working below 

# @api_view(['POST'])
# def login(request):
#     if request.method == 'POST':
#         serializer = UserSerializer(data=request.data)
#         if serializer.is_valid():
#             username = serializer.validated_data['username']
#             password = serializer.validated_data['password']
#             user = authenticate(request, username=username, password=password)
#             if user is not None:
#                 refresh = RefreshToken.for_user(user)
                
#                 # Construct the response with username, email, and name (first_name + last_name)
#                 return Response({
#                     'refresh': str(refresh),
#                     'access': str(refresh.access_token),
#                     'userId': user.id,
#                     'username': user.username,
#                     'email': user.email,
#                     # 'name': f"{user.first_name} {user.last_name}".strip(),  # Combine first and last name
#                 }, status=status.HTTP_200_OK)
#             else:
#                 return Response({'detail': 'Invalid credentials'}, status=status.HTTP_401_UNAUTHORIZED)
#         return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

    # # Handle GET request to fetch superuser details
    # if request.method == 'GET':
    #     superusers = User.objects.filter(is_superuser=True)  # Query for superusers
    #     for user in superusers:
    #       print(user.username, user.email)
    #     print("Superusers:", superusers)  # Debugging print to check data
    #     serializer = UserSerializer(superusers, many=True)
    #     print("Serialized data:", serializer.data)  # Debugging print to check serialized data
    #     return Response(serializer.data, status=status.HTTP_200_OK)



# working below 

# class UserLogoutView(generics.GenericAPIView):
#     serializer_class=LogoutSerializer
#     permission_classes = [IsAuthenticated]
#     def post(self,request,format=None):
#         serializer=self.serializer_class(data=request.data)
#         # https://github.com/jazzband/djangorestframework-simplejwt/issues/218
#         serializer.is_valid(raise_exception=True)
#         # message = serializer.data['refresh_token']
#         # message_bytes = message.encode('ascii')
#         # base64_bytes = base64.b64encode(message_bytes)
#         RefreshToken(serializer.data['refresh_token']).blacklist()

#         return Response({"status":True},status=status.HTTP_200_OK)



class ImportExcelData(APIView):
    permission_classes=[IsAuthenticated,IsSystemOrClientAdmin|IsClientUserEditAndRead]
    # permission_classes=[IsAuthenticated]

    def post(self, request, format=None):
        file_obj = request.FILES.get('file')  # Assuming file is sent in the request
        if not file_obj:
            return Response({"error": config['ERROR_UPLOADING']}, status=status.HTTP_400_BAD_REQUEST)

        company_id = request.data.get('company_id')
        if not company_id:
            return Response({"error": config['ID_REQUIRED']}, status=status.HTTP_400_BAD_REQUEST)
        try:
            # company = Company.objects.get(id=company_id)  # by manish
            company = ClientTemplateCompany.objects.get(id=company_id)

            required_columns = ["Origin Port", "Destination Port", "Transit\ntime", "20'GP", "40'HC", "Effective Date", "Expiration Date"]
            sheets_to_read = ['F.E', 'E.Africa', 'Gulf-Red Sea']  # Adjust sheet names as per your Excel file

            # sheets_to_read = ['E.Africa']  # Adjust sheet names as per your Excel file

            # required_columns = config.get("REQUIRED_COLUMN" , "").split(",")
            # sheets_to_read = config.get("SHEETS_TO_READ" , "").split(",")  # Adjust sheet names as per your Excel file
            # sheets_to_read = ['E.Africa']  # Adjust sheet names as per your Excel file

            combined_df = pd.DataFrame()

            for sheet_name in sheets_to_read:
                df = pd.read_excel(file_obj, sheet_name=sheet_name, header=7, usecols=required_columns)
                df.dropna(how='all', inplace=True)
                df.rename(columns={"Transit\ntime": "Transit time"}, inplace=True)

                # Convert "Destination Port" to uppercase
                # df['Destination Port'] = df['Destination Port'].str.upper().str.replace('PORT', '').str.strip()
                df['Destination Port'] = df['Destination Port'].str.replace('PORT', '').str.strip()

                # Convert Timestamps to string for JSON serialization
                df['Effective Date'] = df['Effective Date'].apply(lambda x: x.strftime('%Y-%m-%d') if isinstance(x, datetime) else None)
                df['Expiration Date'] = df['Expiration Date'].apply(lambda x: x.strftime('%Y-%m-%d') if isinstance(x, datetime) else None)

                # Append the data from the current sheet to the combined DataFrame
                combined_df = pd.concat([combined_df, df], ignore_index=True)

            # Replace NaN or NaT values with None for JSON serialization
            combined_df.replace({np.nan: None, pd.NaT: None}, inplace=True)

            results = combined_df.to_dict(orient='records')
            # Save the data to the database
            freight_types = ["20'GP", "40'HC"]

            for item in results:
                source_name = item["Origin Port"]
                destination_name = item["Destination Port"]
                transit_time_value = item["Transit time"]
                effective_date = item["Effective Date"]
                expiration_date = item["Expiration Date"]

                source, _ = Source.objects.get_or_create(name=source_name)
                destination, _ = Destination.objects.get_or_create(name=destination_name)
                transit_time, _ = TransitTime.objects.get_or_create(time=str(transit_time_value))

                for ft in freight_types:
                    rate_value = item.get(ft)

                    if rate_value is not None:
                        freight_type, _ = FreightType.objects.get_or_create(type=ft)

                        # Check if the rate record already exists
                        existing_rate = Rate.objects.filter(
                            company=company,
                            source=source,
                            destination=destination,
                            freight_type=freight_type
                        ).first()

                        # Fetch all versions sorted by effective_date descending
                        existing_versions = VersionedRate.objects.filter(
                            company=company,
                            source=source,
                            destination=destination,
                            freight_type=freight_type
                        ).order_by('-id')
                        for version in existing_versions:
                            print(f"Existing version found: {version}")
            
                        if existing_rate:
                            # Exclude the current existing_rate's version from existing_versions
                            existing_versions = existing_versions.exclude(id=existing_rate.version.id)

                            # Print existing versions for debugging
                            for version in existing_versions:
                                print(f"Existing version found: {version}")

                            # Select the second latest version for mapping
                            second_latest_version = existing_versions.first() if existing_versions.exists() else None
                            print("second_latest_version:", second_latest_version)

                            # Check if there are changes in rate value, effective date, or expiration date
                            has_changes = (
                                existing_rate.rate != rate_value or
                                existing_rate.effective_date != effective_date or
                                existing_rate.expiration_date != expiration_date
                            )

                            if has_changes:
                                print("Inside has_changes")

                                # Create a new versioned rate
                                VersionedRate.objects.create(
                                    company=company,
                                    source=source,
                                    destination=destination,
                                    transit_time=transit_time,
                                    freight_type=freight_type,
                                    rate=rate_value,
                                    effective_date=effective_date,
                                    expiration_date=expiration_date,
                                    is_current=False
                                )

                                # Update the existing rate
                                existing_rate.rate = rate_value
                                existing_rate.effective_date = effective_date
                                existing_rate.expiration_date = expiration_date
                                if second_latest_version:
                                    existing_rate.version = second_latest_version

                                existing_rate.save()
                        else:
                            print("inside else")
                            # Create new rate and version
                            versioned_rate = VersionedRate.objects.create(
                                company=company,
                                source=source,
                                destination=destination,
                                transit_time=transit_time,
                                freight_type=freight_type,
                                rate=rate_value,
                                effective_date=effective_date,
                                expiration_date=expiration_date,
                                is_current=True
                            )
                            Rate.objects.create(
                                company=company,
                                source=source,
                                destination=destination,
                                transit_time=transit_time,
                                freight_type=freight_type,
                                rate=rate_value,
                                effective_date=effective_date,
                                expiration_date=expiration_date,
                                version=versioned_rate
                            )
            return Response({"message": f"Excel {config['SUCCESS_UPLOADED_MESSAGE']}", "results": results}, status=status.HTTP_201_CREATED)

        except ClientTemplateCompany.DoesNotExist:
            return Response({"error": "Company does not exist"}, status=status.HTTP_400_BAD_REQUEST)
        except Exception as e:
            return Response({"error": str(e)}, status=status.HTTP_400_BAD_REQUEST)
    
class ExtractWordTableView(APIView):
    parser_classes = [MultiPartParser]
    permission_classes=[IsAuthenticated,IsSystemOrClientAdmin|IsClientUserEditAndRead]
    # permission_classes=[IsAuthenticated]

    def post(self, request, *args, **kwargs):
        try:
            # Assume the file is passed as 'file' in the POST request
            file_obj = request.FILES['file']

            # Extract company ID from the request data
            company_id = request.data.get('company_id')
            if not company_id:
                return JsonResponse({"error": config['ID_REQUIRED']}, status=status.HTTP_400_BAD_REQUEST)

            # Process the uploaded Word document
            extracted_data, expiration_date = self.extract_table_data(file_obj)

            # Convert extracted data to desired format
            converted_data = [self.convert_to_desired_format(item, expiration_date) for item in extracted_data]

            # Save the converted data
            self.save_imported_data(converted_data, company_id)

            # Return JSON response
            return Response({"message": f"Word {config['SUCCESS_UPLOADED_MESSAGE']}", "results": converted_data}, status=status.HTTP_201_CREATED)

        except Exception as e:
            # Handle exceptions (e.g., file not found, invalid format, etc.)
            return JsonResponse({"error": str(e)}, status=status.HTTP_400_BAD_REQUEST)

    def extract_table_data(self, file_obj):
        doc = Document(file_obj)
        tables = doc.tables

        # Assuming the table you want is the first table in the document
        table = tables[0]

        data = []
        keys = None

        for i, row in enumerate(table.rows):
            text = [cell.text.strip() for cell in row.cells]
            if i == 0:
                keys = text  # assuming first row as header
            else:
                data.append(dict(zip(keys, text)))
        
        # Extracting the "Valid till" date from the document
        expiration_date = self.extract_valid_till_date(doc)
        
        return data, expiration_date

    def extract_valid_till_date(self, doc):
        valid_till_pattern = re.compile(r"Valid till\s+(\d{2}/\d{2}\s*/\s*\d{4})", re.IGNORECASE)
        
        for paragraph in doc.paragraphs:
            match = valid_till_pattern.search(paragraph.text)
            if match:
                    date_str = match.group(1).replace(" ", "")  # Remove any spaces within the date
                    # Convert the extracted date to a datetime object
                    date_obj = datetime.strptime(date_str, "%d/%m/%Y")
                    # Format the date to YYYY-MM-DD
                    return date_obj.strftime("%Y-%m-%d")  # Remove any spaces within the date
        
        return ""

    def convert_to_desired_format(self, data, expiration_date):
        formatted_data = {
            "Origin Port": "Nhava Sheva",
            "Destination Port": data.get("Sector", ""),
            "Transit time": data.get("Appx.T.time", 0),
            "20'GP": float(data.get("20'GP", 0)),
            "40'HC": float(data.get("40'GP/HC", 0)),
            "Effective Date": "2024-04-01",
            "Expiration Date": expiration_date if expiration_date else data.get("Expiration Date", "")
        }
        return formatted_data

    def save_imported_data(self, results, company_id):
        freight_types = ["20'GP", "40'HC"]  # Define the freight types you are processing

        for item in results:
            source_name = item["Origin Port"]
            destination_name = item["Destination Port"]
            transit_time_value = item["Transit time"]
            effective_date = item["Effective Date"]
            expiration_date = item["Expiration Date"]
            print(f"Source: {source_name}, Destination: {destination_name}, Transit Time: {transit_time_value}, "
                f"Effective Date: {effective_date}, Expiration Date: {expiration_date}")
            source, _ = Source.objects.get_or_create(name=source_name)
            destination, _ = Destination.objects.get_or_create(name=destination_name)
            transit_time, _ = TransitTime.objects.get_or_create(time=str(transit_time_value))

            for ft in freight_types:
                rate_value = item.get(ft)
                if rate_value is not None:
                    # company = get_object_or_404(Company, id=company_id) # by manish
                    company = get_object_or_404(ClientTemplateCompany, id=company_id)

                    freight_type, _ = FreightType.objects.get_or_create(type=ft)

                    # Check if the rate record already exists
                    existing_rate = Rate.objects.filter(
                        company=company,
                        source=source,
                        destination=destination,
                        freight_type=freight_type
                    ).first()

                    # Fetch all versions sorted by effective_date descending
                    existing_versions = VersionedRate.objects.filter(
                        company=company,
                        source=source,
                        destination=destination,
                        freight_type=freight_type
                    ).order_by('-id')
                    # Exclude the current existing_rate
                    if existing_rate:
                        existing_versions = existing_versions.exclude(id=existing_rate.version.id)
                        for version in existing_versions:
                            print(f"Existing version found: {version}")

                    # Select the second latest version for mapping
                    second_latest_version = existing_versions[0] if len(existing_versions) > 0 else None
                    print("second_latest_version,second_latest_version",  second_latest_version)

                    if existing_rate:
                        has_changes = (
                            existing_rate.rate != rate_value or
                            existing_rate.effective_date != effective_date or
                            existing_rate.expiration_date != expiration_date
                        )

                        if has_changes:
                            VersionedRate.objects.create(
                                company=company,
                                source=source,
                                destination=destination,
                                transit_time=transit_time,
                                freight_type=freight_type,
                                rate=rate_value,
                                effective_date=effective_date,
                                expiration_date=expiration_date,
                                is_current=False
                            )
                            # Update the existing rate
                            existing_rate.rate = rate_value
                            existing_rate.effective_date = effective_date
                            existing_rate.expiration_date = expiration_date
                            if second_latest_version:
                                existing_rate.version = second_latest_version

                            existing_rate.save()
                    else:
                        # Create new rate and version
                        versioned_rate = VersionedRate.objects.create(
                            company=company,
                            source=source,
                            destination=destination,
                            transit_time=transit_time,
                            freight_type=freight_type,
                            rate=rate_value,
                            effective_date=effective_date,
                            expiration_date=expiration_date,
                            is_current=True
                        )
                        Rate.objects.create(
                            company=company,
                            source=source,
                            destination=destination,
                            transit_time=transit_time,
                            freight_type=freight_type,
                            rate=rate_value,
                            effective_date=effective_date,
                            expiration_date=expiration_date,
                            version=versioned_rate
                        )
                    # if existing_rate:
                    #     # Exclude the current existing_rate's version from existing_versions
                    #     existing_versions = existing_versions.exclude(id=existing_rate.version.id)

                    #     # Print existing versions for debugging
                    #     for version in existing_versions:
                    #         print(f"Existing version found: {version}")

                    #     # Select the second latest version for mapping
                    #     second_latest_version = existing_versions.first() if existing_versions.exists() else None
                    #     print("second_latest_version:", second_latest_version)

                    #     # Check if there are changes in rate value, effective date, or expiration date
                    #     has_changes = (
                    #         existing_rate.rate != rate_value or
                    #         existing_rate.effective_date != effective_date or
                    #         existing_rate.expiration_date != expiration_date
                    #     )

                    #     if has_changes:
                    #         print("Inside has_changes")

                    #         # Create a new versioned rate
                    #         VersionedRate.objects.create(
                    #             company=company,
                    #             source=source,
                    #             destination=destination,
                    #             transit_time=transit_time,
                    #             freight_type=freight_type,
                    #             rate=rate_value,
                    #             effective_date=effective_date,
                    #             expiration_date=expiration_date,
                    #             is_current=False
                    #         )

                    #         # Update the existing rate
                    #         existing_rate.rate = rate_value
                    #         existing_rate.effective_date = effective_date
                    #         existing_rate.expiration_date = expiration_date
                    #         if second_latest_version:
                    #             existing_rate.version = second_latest_version

                    #         existing_rate.save()
                    # else:
                    #     print("inside else")
                    #     # Create new rate and version
                    #     versioned_rate = VersionedRate.objects.create(
                    #         company=company,
                    #         source=source,
                    #         destination=destination,
                    #         transit_time=transit_time,
                    #         freight_type=freight_type,
                    #         rate=rate_value,
                    #         effective_date=effective_date,
                    #         expiration_date=expiration_date,
                    #         is_current=True
                    #     )
                    #     Rate.objects.create(
                    #         company=company,
                    #         source=source,
                    #         destination=destination,
                    #         transit_time=transit_time,
                    #         freight_type=freight_type,
                    #         rate=rate_value,
                    #         effective_date=effective_date,
                    #         expiration_date=expiration_date,
                    #         version=versioned_rate
                    #         )

class ExtractPDFTableView(APIView):
    permission_classes=[IsAuthenticated,IsSystemOrClientAdmin|IsClientUserEditAndRead]
    # permission_classes=[IsAuthenticated]

 
    def extract_valid_dates(self, text):
        valid_dates = {}

        # Regex pattern to find valid dates and regions
        pattern = r"valid(?:\s+FROM)?\s+(\d{2}-\d{2}-\d{4})\s+to\s+(\d{2}-\d{2}-\d{4})\s+(Vessel Sailing\s+)?([A-Z\s-]+)"
        matches = re.finditer(pattern, text, re.IGNORECASE)

        for match in matches:
            start_date = datetime.strptime(match.group(1), "%d-%m-%Y").date()
            end_date = datetime.strptime(match.group(2), "%d-%m-%Y").date()
            region = match.group(4).strip()

            if '\n' in region:
                region = region.split('\n', 1)[0].strip()

            # Remove "Vessel Sailing" prefix if present
            if region.startswith("Vessel Sailing"):
                region = region[len("Vessel Sailing"):].strip()

            valid_dates[region] = {"valid_from": start_date, "valid_to": end_date}


        return valid_dates
    
    def process_document(self, content, mime_type):
        KEYFILE_PATH = r'E:\PROJECT\MooMove\app\shipment\aggregator\jspl-trocr-2024-419309-21a5c31d8f62.json'
        # Replace these with your actual Google Cloud project ID and processor ID
        project_id = config['PROJECT_ID']
        location = config['LOCATION']
        processor_id = config['PROCESSOR_ID']

        # Authenticate with service account JSON key file
        credentials = service_account.Credentials.from_service_account_file(KEYFILE_PATH)
        client = documentai.DocumentProcessorServiceClient(credentials=credentials)

        # The full resource name of the processor
        name = f'projects/{project_id}/locations/{location}/processors/{processor_id}'

        # Configure the process request
        request = {"name": name, "raw_document": {"content": content, "mime_type": mime_type}}

        # Recognize text entities in the document
        result = client.process_document(request=request)

        document = result.document

        # Get the document text
        document_text = document.text
        # print("document_text: ",document_text)
        # Get individual entities
        entities_text = ""
        for entity in document.entities:
            entities_text += f"<li>Entity type: {entity.type_}, Value: {entity.mention_text}</li>"

        return document_text, entities_text
    
    def clean_region_name(self, region_name):
        # Clean up region name to match keys in valid_dates
        return region_name.replace('\n', '').strip()


    def apply_valid_dates(self, tables, valid_dates):
        for table in tables:
            region = self.clean_region_name(table.get('Region', ''))
            

            if region in valid_dates:
                table['Valid From'] = valid_dates[region]['valid_from']
                table['Valid To'] = valid_dates[region]['valid_to']
            elif region  == "WEST AFRICA":
                # Assign same dates as SOUTH WEST AFRICA
                south_west_africa_dates = valid_dates.get('SOUTH WEST AFRICA')
                if south_west_africa_dates:
                    table['Valid From'] = south_west_africa_dates['valid_from']
                    table['Valid To'] = south_west_africa_dates['valid_to']
            else:
                table['Valid From'] = None
                table['Valid To'] = None

    def post(self, request):

        try:
            # Check if the 'file' key is present in request.FILES
            if 'file' not in request.FILES:
                return Response({"error": "No 'file' key found in request data"}, status=status.HTTP_400_BAD_REQUEST)

            # Extract company_id from request data or session if authenticated
            company_id = request.data.get('company_id')  # Adjust this based on how company_id is passed

            # Retrieve company instance based on company_id
            # company = Company.objects.get(id=company_id)  # Assuming Company model and id field exist   # by manish
            company = ClientTemplateCompany.objects.get(id=company_id)  # Assuming Company model and id field exist

            pdf_file = request.FILES['file']
            content = pdf_file.read()
            mime_type = pdf_file.content_type

            # Process the document
            document_text, entities_text = self.process_document(content, mime_type)
            all_tables = []
            regions = {
                "WEST AFRICA": ["LOME", "TEMA", "ABIDJAN", "COTONOU", "TINCAN ISLAND", "LAGOS", "ONNE", "DAKAR"],
                "SOUTH WEST AFRICA": ["LUANDA", "POINTE NOIRE", "DOUALA", "MATADI", "WALVIS BAY", "KRIBI"],
                "EAST AFRICA": ["MOMBASA", "DAR-ES-SALAAM"],
                "SOUTH AFRICA": ["DURBAN", "CAPETOWN"]
            }

            reader = PyPDF2.PdfReader(pdf_file)
            num_pages = len(reader.pages)

            for page_num in range(num_pages):
                page = reader.pages[page_num]
                text = page.extract_text()
                # print("dockumet text: ", document_text)
                # print("text: ", text)
                # Extract valid dates for each region
                valid_dates = self.extract_valid_dates(document_text)
                

                # Regex pattern for main table rows like 'MATADI 2075 3450 2075 3450'
                table_pattern = r"([A-Z\s-]+)\s+(\d+)\s+(\d+)\s+(\d+)\s+(\d+)"
                table_matches = re.finditer(table_pattern, text, re.MULTILINE)

                # Iterate through each match and capture row number
                for idx, match in enumerate(table_matches, start=1):
                    row_number = match.start()
                    # port_name = match.group(1).strip().replace("PORT", "").strip()  # Remove "PORT"
                    # port_name = match.group(1).split('\n', 1)[-1].strip().replace("PORT", "").strip()
                    port_name = match.group(1).split('\n', 1)[-1].strip().replace("PORT", "").strip().replace("  ", " ").replace(" -", "-")


                    rates_ex_nhava_sheva = {
                        "origin_port": "Nhava Sheva",
                        "20'GP": int(match.group(2)),
                        "40'HC": int(match.group(3))
                    }
                    rates_ex_mundra = {
                        "origin_port": "Mundra",
                        "20'GP": int(match.group(4)),
                        "40'HC": int(match.group(5))
                    }

                    # Determine region based on destination port
                    region = None
                    for key, value in regions.items():
                        if any(port in port_name for port in value):
                            region = key
                            break

                    table_data = {
                        "Row Number": row_number,
                        "Destination Port": port_name,
                        "Origin Port Nhava Sheva": rates_ex_nhava_sheva,
                        "Origin Port Mundra": rates_ex_mundra,
                        "Transit time": None,  # Initialize as None
                        "Region": region,  # Assign region
                        "Valid From": valid_dates[region]["valid_from"] if region in valid_dates else None,
                        "Valid To": valid_dates[region]["valid_to"] if region in valid_dates else None
                    }
                     
                    all_tables.append(table_data)
                self.apply_valid_dates(all_tables, valid_dates)
                # Regex pattern to find transit time like '42 days'
                days_pattern = r"(\d+)\s+days"
                days_matches = re.finditer(days_pattern, text)

                # Find the corresponding transit time and update in all_tables
                for days_match in days_matches:
                    transit_time = int(days_match.group(1))

                    # Find the nearest row number before the transit time match
                    row_number = -1
                    for table_data in all_tables:
                        if table_data["Row Number"] < days_match.start():
                            row_number = table_data["Row Number"]
                        else:
                            break

                    # Update transit time in the corresponding row
                    if row_number != -1:
                        for table_data in all_tables:
                            if table_data["Row Number"] == row_number:
                                table_data["Transit time"] = transit_time
                                break
            if all_tables:
                created_records = []

                for item in all_tables:
                    destination_port = item['Destination Port']
                    transit_time = item['Transit time']
                    region = item['Region']
                    valid_from = item['Valid From']
                    valid_to = item['Valid To']

                    origin_ports = [item['Origin Port Nhava Sheva'], item['Origin Port Mundra']]

                    for origin_port in origin_ports:
                        record = {
                            "Destination Port": destination_port,
                            "Origin Port": origin_port['origin_port'] if 'origin_port' in origin_port else origin_port['origin_port'],
                            "20'GP": origin_port["20'GP"],
                            "40'HC": origin_port["40'HC"],
                            "Transit time": transit_time,
                            "effective_date": valid_from,
                            "expiration_date": valid_to
                        }

                        created_records.append(record)

                # Save created_records to database
                for item in created_records:
                    source_name = item["Origin Port"]
                    destination_name = item["Destination Port"]
                    transit_time_value = item["Transit time"]
                    effective_date = item["effective_date"]
                    expiration_date = item["expiration_date"]

                    source, _ = Source.objects.get_or_create(name=source_name)
                    destination, _ = Destination.objects.get_or_create(name=destination_name)
                    transit_time, _ = TransitTime.objects.get_or_create(time=str(transit_time_value))

                    freight_types = ["20'GP", "40'HC"]
                    for ft in freight_types:
                        rate_value = item.get(ft)

                        if rate_value is not None:
                            freight_type, _ = FreightType.objects.get_or_create(type=ft)

                            # Check if the rate record already exists
                            existing_rate = Rate.objects.filter(
                                company=company,
                                source=source,
                                destination=destination,
                                freight_type=freight_type
                            ).first()

                            # Fetch all versions sorted by effective_date descending
                            existing_versions = VersionedRate.objects.filter(
                                company=company,
                                source=source,
                                destination=destination,
                                freight_type=freight_type
                            ).order_by('-id')
                            for version in existing_versions:
                                print(f"Existing version found1: {version}")

                            if existing_rate:
                                # Exclude the current existing_rate's version from existing_versions
                                existing_versions = existing_versions.exclude(id=existing_rate.version.id)

                                # Print existing versions for debugging
                                for version in existing_versions:
                                    print(f"Existing version found: {version}")

                                # Select the second latest version for mapping
                                second_latest_version = existing_versions.first() if existing_versions.exists() else None
                                print("second_latest_version:", second_latest_version)

                                # Check if there are changes in rate value, effective date, or expiration date
                                has_changes = (
                                    existing_rate.rate != rate_value or
                                    existing_rate.effective_date != effective_date or
                                    existing_rate.expiration_date != expiration_date
                                )

                                if has_changes:
                                    print("Inside has_changes")

                                    # Create a new versioned rate
                                    VersionedRate.objects.create(
                                        company=company,
                                        source=source,
                                        destination=destination,
                                        transit_time=transit_time,
                                        freight_type=freight_type,
                                        rate=rate_value,
                                        effective_date=effective_date,
                                        expiration_date=expiration_date,
                                        is_current=False
                                    )

                                    # Update the existing rate
                                    existing_rate.rate = rate_value
                                    existing_rate.effective_date = effective_date
                                    existing_rate.expiration_date = expiration_date
                                    if second_latest_version:
                                        existing_rate.version = second_latest_version

                                    existing_rate.save()
                            else:
                                # Create new rate and version
                                versioned_rate = VersionedRate.objects.create(
                                    company=company,
                                    source=source,
                                    destination=destination,
                                    transit_time=transit_time,
                                    freight_type=freight_type,
                                    rate=rate_value,
                                    effective_date=effective_date,
                                    expiration_date=expiration_date,
                                    is_current=True
                                )
                                Rate.objects.create(
                                    company=company,
                                    source=source,
                                    destination=destination,
                                    transit_time=transit_time,
                                    freight_type=freight_type,
                                    rate=rate_value,
                                    effective_date=effective_date,
                                    expiration_date=expiration_date,
                                    version=versioned_rate
                                )

                return Response({"message": f"PDF {config['SUCCESS_UPLOADED_MESSAGE']}", "tables": all_tables}, status=status.HTTP_200_OK)

        except KeyError:
            return Response({"error": "No 'file' key found in request data"}, status=status.HTTP_400_BAD_REQUEST)

        except Exception as e:
            return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

class RateWithVersionsAPIView(APIView):
    permission_classes=[IsAuthenticated,IsClientUserEditAndRead|IsSystemOrClientAdmin|IsClientUserReadOnly|IsUser]
    # permission_classes=[IsAuthenticated]

    def get(self, request, company_id):
        # print("company_id: ", company_id)
        try:
            # Fetch the latest rates (is_current=True) for the given company
            rates = Rate.objects.filter(company_id=company_id,soft_delete=False)
            # rates = Rate.objects.filter(soft_delete=False)
            
            # rates = Rate.objects.filter(company_id=company_id, version__is_current=True)
            # if not rates:
            #     print('rates form 807')
            #     rates = Rate.objects.filter(company_id=company_id, version__is_current=False)
            # Serialize the rates and their versioned rates
            # print(Rate.objects.filter(company_id=company_id).count(),
            #       Rate.objects.filter(company_id=company_id, version__is_current=False).count())

            # rate_serializer = RateSerializer(rates, many=True)
            rate_serializer = RateSerializer(rates, many=True)
            # print(rate_serializer.data)
            # count = len(rate_serializer.data)
            # print("count: ", count)
            return Response(rate_serializer.data, status=status.HTTP_200_OK)

        except Rate.DoesNotExist:
            return Response({"error": config['RATE_DATA_NOT_EXIST']}, status=status.HTTP_404_NOT_FOUND)
        except Exception as e:
            return Response({"error": str(e)}, status=status.HTTP_400_BAD_REQUEST)
   
class CompanyListAPIView(APIView):
    permission_classes=[IsAuthenticated,IsClientUserEditAndRead|IsSystemOrClientAdmin|IsClientUserReadOnly|IsUser]
    # permission_classes=[IsAuthenticated]

    def get(self, request):
        companies = Company.objects.filter(soft_delete=False)
        serializer = CompanySerializer(companies, many=True)
        return Response(serializer.data)

    def post(self, request):
        serializer = CompanySerializer(data=request.data)
        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data, status=status.HTTP_201_CREATED)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
    
class ClientTemplateCompanyAPIView(APIView):
    permission_classes=[IsAuthenticated,IsClientUserEditAndRead|IsSystemOrClientAdmin|IsClientUserReadOnly|IsUser]
    # permission_classes=[IsAuthenticated]

    def get(self, request):
        companies = ClientTemplateCompany.objects.filter(soft_delete=False)
        serializer = ClientTemplateCompanySerializer(companies, many=True)
        return Response(serializer.data)

    def post(self, request):
        serializer = ClientTemplateCompanySerializer(data=request.data)
        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data, status=status.HTTP_201_CREATED)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
    
class SourceListAPIView(APIView):
    permission_classes=[IsAuthenticated,IsClientUserEditAndRead|IsSystemOrClientAdmin|IsClientUserReadOnly|IsUser]
    # permission_classes=[IsAuthenticated]

    def get(self, request):
        sources = Source.objects.filter(soft_delete=False)
        serializer = SourceSerializer(sources, many=True)
        return Response(serializer.data)

    # def post(self, request):
    #     serializer = CompanySerializer(data=request.data)
    #     if serializer.is_valid():
    #         serializer.save()
    #         return Response(serializer.data, status=status.HTTP_201_CREATED)
    #     return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


    def post(self, request):
        try:
            with transaction.atomic():
                requestData = request.data
                source_name = requestData.get('name')

                # Check if a source with the same name already exists
                existing_source = Source.objects.filter(name=source_name).first()
                if existing_source:
                    return Response({"message": f"'{source_name}' source already exists"}, status=status.HTTP_200_OK)

                # Create a new source if it doesn't exist
                Source.objects.create(name=source_name)
                return Response({'message': f"'{source_name}' source successfully created"}, status=status.HTTP_201_CREATED)

        except Exception as e:
            return Response({'message': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

class DestinaltionListAPIView(APIView):
    permission_classes=[IsAuthenticated,IsClientUserEditAndRead|IsSystemOrClientAdmin|IsClientUserReadOnly|IsUser]
    # permission_classes=[IsAuthenticated]

    def get(self, request):
        destination = Destination.objects.filter(soft_delete=False)
        serializer = DestinationSerializer(destination, many=True)
        return Response(serializer.data)
    

    def post(self, request):
        try:
            with transaction.atomic():
                requestData = request.data
                destination_name = requestData.get('name')

                # Check if a destination with the same name already exists
                existing_source = Destination.objects.filter(name=destination_name).first()
                if existing_source:
                    return Response({"message": f"'{destination_name}' destination already exists"}, status=status.HTTP_200_OK)

                # Create a new destination if it doesn't exist
                Destination.objects.create(name=destination_name)
                return Response({'message': f"'{destination_name}' destination successfully created"}, status=status.HTTP_201_CREATED)

        except Exception as e:
            return Response({'message': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
    
class FrightTypeListAPIView(APIView):
    permission_classes=[IsAuthenticated,IsClientUserEditAndRead|IsSystemOrClientAdmin|IsClientUserReadOnly|IsUser]
    # permission_classes=[IsAuthenticated]

    def get(self, request):
        freight_type = FreightType.objects.all()
        serializer = FreightTypeSerializer(freight_type, many=True)
        return Response(serializer.data)

    def post(self, request):
        try:
           requestData = request.data
           freight_type = requestData.get('type')
           # Check if a FreightType with this type already exists
           existing_freight_type = FreightType.objects.filter(type=freight_type).first()
           if existing_freight_type:
                # If the type already exists, return the message
                return Response({"message": config['EXIST_DATA']}, status=status.HTTP_200_OK)
           
            # Otherwise, create a new FreightType instance
           FreightType.objects.create(type=freight_type)

           return Response({'message': config['FREIGHT_TYPE_CREATED']}, status=status.HTTP_201_CREATED)


        except Exception as e:
           return Response({'detail': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)    

class RateListView(generics.ListAPIView):
    permission_classes =[IsAuthenticated,IsClientUserEditAndRead|IsSystemOrClientAdmin|IsClientUserReadOnly|IsUser]
    # permission_classes = [IsAuthenticated]
    serializer_class = RateSerializer1

    def get_queryset(self):
        source_id = self.kwargs.get('source')
        destination_id = self.kwargs.get('destination')

        with connection.cursor() as cursor:
            # THIS IS STORE PRODUCER NAME IS "get_combined_rates_data" USE HERE 
            cursor.execute("SELECT * FROM get_combined_rates_data(%s, %s)", [source_id, destination_id])
            rows = cursor.fetchall()

            # columns = [
            #     'id', 'unique_uuid', 'company_id', 'company_name', 'rate', 'currency',
            #     'free_days', 'spot_filed', 'transhipment_add_port', 'effective_date',
            #     'expiration_date', 'un_number' ,'vessel_name', 'cargotype', 'voyage', 'hazardous', 'terms_condition', 'source_id', 'source_name', 'destination_id', 
            #     'destination_name','transit_time_id','transit_time','freight_type_id','freight_type',
            #     # 'client_template_id', 'client_template_name',
            # ]
            columns = config.get("RATE_LIST_QUERYSET" , "").split(",")
            data = [dict(zip(columns, row)) for row in rows]
            return data  # Return as list of dictionaries for serialization

    
# class RateListView(generics.ListAPIView):
    # permission_classes=[IsAuthenticated]

    # serializer_class = RateSerializer1

    # def get_queryset(self):
    #     # Extract parameters from URL path variables
    #     source_id = self.kwargs.get('source')
    #     destination_id = self.kwargs.get('destination')
    #     # freight_type_id = self.kwargs.get('freight_type')

    #      # Call the stored procedure and fetch results
    #     with connection.cursor() as cursor:
    #         cursor.execute("SELECT * FROM get_combined_rates_data2(%s, %s)", [source_id, destination_id])
    #         # Fetch all results
    #         rows = cursor.fetchall()
    #        # Manually map tuples to dictionaries
    #         columns = [
    #                 'id', 'company_id','rate', 'currency', 
    #                 'free_days', 'spot_filed', 'transhipment_add_port', 
    #                 'effective_date', 'vessel_name', 'source_id', 'destination_id'
    #             ]
    #         data = [dict(zip(columns, row)) for row in rows]
    #         print(data)
            
      
    #     return data    

        # # Query Rate objects based on parameters
        # rate_queryset = Rate.objects.filter(
        #     source_id=source_id,
        #     destination_id=destination_id,
        #     soft_delete=False
        #     # freight_type_id=freight_type_id
        # )

        # return rate_queryset





class ManualRateWithRateWithVersionsAPIView(APIView):
    permission_classes=[IsAuthenticated,IsClientUserEditAndRead|IsSystemOrClientAdmin|IsClientUserReadOnly]
    # permission_classes=[IsAuthenticated]

    def get(self, request, company_id):
        # print("clienttemplatecompany_id: ", clienttemplatecompany_id)
        try:
            manual_rates = ManualRate.objects.filter(company_id=company_id,soft_delete=False)
            manual_rates_serializer = ManualRateSerializer(manual_rates, many=True)
            return Response(manual_rates_serializer.data, status=status.HTTP_200_OK)

        except ManualRate.DoesNotExist:
            return Response({"error": "ManualRate version not found for the company"}, status=status.HTTP_404_NOT_FOUND)
        except Exception as e:
            return Response({"error": str(e)}, status=status.HTTP_400_BAD_REQUEST)

class ManualRateListView(APIView):
    permission_classes = [IsAuthenticated|IsClientUserEditAndRead|IsSystemAdministrator|IsClientAdministrator]
    # permission_classes = [IsAuthenticated]

    #  FOR GET 
    def get(self, request):
        manual_rate = ManualRate.objects.filter(soft_delete=False)  # Exclude soft-deleted records
        manual_rate_serializer = ManualRateSerializer(manual_rate, many=True)
        return Response(manual_rate_serializer.data)

    #  FOR POST 

    def post(self, request):
        try:
            with transaction.atomic():   
                requestData = request.data
                print(requestData)

                # Extract request data
                company_name = requestData.get('company')
                source_name = requestData.get('source')
                destination_name = requestData.get('destination')
                freight_type = requestData.get('freight_type')
                transit_time = requestData.get('transit_time')
                commodity_name = requestData.get('cargotype')
                manualrate = requestData.get('rate')

                # Ensure required fields are provided
                required_fields = [company_name, source_name, destination_name, freight_type, transit_time,manualrate]
                if not all(required_fields):
                    return Response({"message": required_fields }, status=status.HTTP_400_BAD_REQUEST)

                # Create or retrieve the instances
                # company_instance, _ = Company.objects.get_or_create(name=company_name)
                company_instance, _ = Company.objects.get_or_create(name=company_name)
                # company_instance, _ = ManualShippingList.objects.get_or_create(name=company_name)
                source_instance, _ = Source.objects.get_or_create(name=source_name)
                destination_instance, _ = Destination.objects.get_or_create(name=destination_name)
                freight_type_instance, _ = FreightType.objects.get_or_create(type=freight_type)
                transit_time_instance, _ = TransitTime.objects.get_or_create(time=transit_time)
                commodity_name_instance, _ = Comodity.objects.get_or_create(name=commodity_name)

                # Check for existing records across ManualRate, VersionedRate, and Rate
                filters = {
                    'company': company_instance,
                    'source': source_instance,
                    'destination': destination_instance,
                    'freight_type': freight_type_instance,
                    'transit_time': transit_time_instance,
                    'cargotype': commodity_name_instance,
                    'rate': manualrate,
                    'soft_delete': False  # Exclude soft-deleted records
                }

                existing_manual_rate = ManualRate.objects.filter(**filters).first()
                # existing_versioned_rate = VersionedRate.objects.filter(**filters).first()
                # existing_rate = Rate.objects.filter(**filters).first()

                # If all values are the same in all tables, return a message saying 'Data already exists'
                if existing_manual_rate:
                    return Response({"message": "already exists"}, status=status.HTTP_200_OK)

                # Get the current timestamp in a unique format
                timestamp = datetime.now().strftime('%Y%m%d%H%M%S%f')
                # Generate a UUID and replace hyphens
                unique_id = f"{timestamp}{str(uuid.uuid4()).replace('-', '')[:8]}"
                common_uuid = unique_id[:24]

                # # Create VersionedRate record
                # versioned_rate = VersionedRate.objects.create(
                #     unique_uuid=common_uuid,
                #     company=company_instance,
                #     source=source_instance,
                #     destination=destination_instance,
                #     freight_type=freight_type_instance,
                #     transit_time=transit_time_instance,
                #     cargotype=commodity_name_instance,
                #     rate=requestData.get('rate'),
                #     free_days=requestData.get('free_days'),
                #     free_days_comment=requestData.get('free_days_comment'),
                #     currency=requestData.get('currency'),
                #     hazardous=requestData.get('hazardous'),
                #     un_number=requestData.get('un_number'),
                #     spot_filed=requestData.get('spot_filed'),
                #     isRateTypeStatus=requestData.get('isRateTypeStatus'),
                #     vessel_name=requestData.get('vessel_name'),
                #     voyage=requestData.get('voyage'),
                #     haz_class=requestData.get('haz_class'),
                #     packing_group=requestData.get('packing_group'),
                #     transhipment_add_port=requestData.get('transhipment_add_port'),
                #     effective_date=requestData.get('effective_date'),
                #     expiration_date=requestData.get('expiration_date'),
                #     remarks=requestData.get('remarks'),
                #     terms_condition=requestData.get('terms_condition'),
                #     is_current=True
                # )

                # Create ManualRate record
                manual_rate = ManualRate.objects.create(
                    unique_uuid=common_uuid,
                    company=company_instance,
                    source=source_instance,
                    destination=destination_instance,
                    freight_type=freight_type_instance,
                    transit_time=transit_time_instance,
                    cargotype=commodity_name_instance,
                    rate=manualrate,
                    free_days=requestData.get('free_days'),
                    free_days_comment=requestData.get('free_days_comment'),
                    currency=requestData.get('currency'),
                    hazardous=requestData.get('hazardous'),
                    un_number=requestData.get('un_number'),
                    direct_shipment=requestData.get('direct_shipment'),
                    spot_filed=requestData.get('spot_filed'),
                    isRateTypeStatus=requestData.get('isRateTypeStatus'),
                    vessel_name=requestData.get('vessel_name'),
                    voyage=requestData.get('voyage'),
                    haz_class=requestData.get('haz_class'),
                    packing_group=requestData.get('packing_group'),
                    transhipment_add_port=requestData.get('transhipment_add_port'),
                    effective_date=requestData.get('effective_date'),
                    expiration_date=requestData.get('expiration_date'),
                    remarks=requestData.get('remarks'),
                    terms_condition=requestData.get('terms_condition'),
                    # version=versioned_rate
                )

                # # Create Rate record
                # Rate.objects.create(
                #     unique_uuid=common_uuid,
                #     company=company_instance,
                #     source=source_instance,
                #     destination=destination_instance,
                #     freight_type=freight_type_instance,
                #     transit_time=transit_time_instance,
                #     cargotype=commodity_name_instance,
                #     rate=requestData.get('rate'),
                #     free_days=requestData.get('free_days'),
                #     free_days_comment=requestData.get('free_days_comment'),
                #     currency=requestData.get('currency'),
                #     hazardous=requestData.get('hazardous'),
                #     un_number=requestData.get('un_number'),
                #     spot_filed=requestData.get('spot_filed'),
                #     isRateTypeStatus=requestData.get('isRateTypeStatus'),
                #     vessel_name=requestData.get('vessel_name'),
                #     voyage=requestData.get('voyage'),
                #     haz_class=requestData.get('haz_class'),
                #     packing_group=requestData.get('packing_group'),
                #     transhipment_add_port=requestData.get('transhipment_add_port'),
                #     effective_date=requestData.get('effective_date'),
                #     expiration_date=requestData.get('expiration_date'),
                #     version=versioned_rate,
                #     remarks=requestData.get('remarks'),
                #     terms_condition=requestData.get('terms_condition'),
                     
                # )

                return Response({'message': 'Manual rate processed successfully'}, status=status.HTTP_201_CREATED)

        except Exception as e:
            return Response({'detail': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


    # UPDATING FUNCTION HERE
    
    def put(self, request, unique_uuid):
        try:
            with transaction.atomic():
                requestData = request.data

                # Retrieve the existing ManualRate object
                try:
                    manual_rate_instance = ManualRate.objects.get(unique_uuid=unique_uuid)
                except ManualRate.DoesNotExist:
                    return Response({"detail": "ManualRate not found."}, status=status.HTTP_404_NOT_FOUND)

                # Fetch related instances
                try:
                    company_instance = Company.objects.get(name=requestData.get('company'))
                    source_instance = Source.objects.get(name=requestData.get('source'))
                    destination_instance = Destination.objects.get(name=requestData.get('destination'))
                    freight_type_instance = FreightType.objects.get(type=requestData.get('freight_type'))
                    transit_time_instance = TransitTime.objects.get(time=requestData.get('transit_time'))
                except (Company.DoesNotExist, Source.DoesNotExist, Destination.DoesNotExist, FreightType.DoesNotExist, TransitTime.DoesNotExist) as e:
                    return Response({"detail": str(e).split('DoesNotExist')[0] + " not found."}, status=status.HTTP_404_NOT_FOUND)

                # Update fields in the ManualRate instance
                manual_rate_instance.company = company_instance
                manual_rate_instance.source = source_instance
                manual_rate_instance.destination = destination_instance
                manual_rate_instance.freight_type = freight_type_instance
                manual_rate_instance.transit_time = transit_time_instance

                # Update optional fields with fallback to current values
                manual_rate_instance.cargotype = requestData.get('cargotype', manual_rate_instance.cargotype)
                manual_rate_instance.rate = float(requestData.get('rate', manual_rate_instance.rate))
                manual_rate_instance.direct_shipment = requestData.get('direct_shipment', manual_rate_instance.direct_shipment)
                manual_rate_instance.spot_filed = requestData.get('spot_filed', manual_rate_instance.spot_filed)
                manual_rate_instance.isRateTypeStatus = requestData.get('isRateTypeStatus', manual_rate_instance.isRateTypeStatus)
                manual_rate_instance.hazardous = requestData.get('hazardous', manual_rate_instance.hazardous)
                manual_rate_instance.un_number = requestData.get('un_number', manual_rate_instance.un_number)
                manual_rate_instance.currency = requestData.get('currency', manual_rate_instance.currency)
                manual_rate_instance.vessel_name = requestData.get('vessel_name', manual_rate_instance.vessel_name)
                manual_rate_instance.voyage = requestData.get('voyage', manual_rate_instance.voyage)
                manual_rate_instance.haz_class = requestData.get('haz_class', manual_rate_instance.haz_class)
                manual_rate_instance.packing_group = requestData.get('packing_group', manual_rate_instance.packing_group)
                
                # Convert free_days to integer
                free_days_value = requestData.get('free_days', manual_rate_instance.free_days)
                manual_rate_instance.free_days = int(free_days_value) if free_days_value is not None else manual_rate_instance.free_days
                
                # Update additional fields
                manual_rate_instance.free_days_comment = requestData.get('free_days_comment', manual_rate_instance.free_days_comment)
                manual_rate_instance.transhipment_add_port = requestData.get('transhipment_add_port', manual_rate_instance.transhipment_add_port)
                
                # Handle date fields
                manual_rate_instance.effective_date = requestData.get('effective_date', manual_rate_instance.effective_date)
                manual_rate_instance.expiration_date = requestData.get('expiration_date', manual_rate_instance.expiration_date)
                
                manual_rate_instance.remarks = requestData.get('remarks', manual_rate_instance.remarks)
                manual_rate_instance.terms_condition = requestData.get('terms_condition', manual_rate_instance.terms_condition)
                manual_rate_instance.isRateUsed = requestData.get('isRateUsed', manual_rate_instance.isRateUsed)

                # Save changes
                manual_rate_instance.save()
                
                # Return success response
                return Response({"detail": "ManualRate updated successfully."}, status=status.HTTP_200_OK)

        except Exception as e:
            return Response({"detail": str(e)}, status=status.HTTP_400_BAD_REQUEST)

   

    #  FOR DETELE

    def delete(self, request, unique_uuid):
        try:
        # Start a transaction to ensure atomic updates
            with transaction.atomic():
                # Retrieve all related records with the same unique_uuid
                manual_rate_instance = ManualRate.objects.get(unique_uuid=unique_uuid, soft_delete=False)
                # version_rate_instance = VersionedRate.objects.get(unique_uuid=unique_uuid, soft_delete=False)
                # rate_instance = Rate.objects.get(unique_uuid=unique_uuid, soft_delete=False)
                # company_instance = Company.objects.get(unique_uuid=unique_uuid, soft_delete=False)
                # source_instance = Source.objects.get(unique_uuid=unique_uuid, soft_delete=False)
                # destination_instance = Destination.objects.get(unique_uuid=unique_uuid, soft_delete=False)

                # Perform soft deletion for all records
                manual_rate_instance.soft_delete = True
                # version_rate_instance.soft_delete = True
                # rate_instance.soft_delete = True
                # company_instance.soft_delete = True
                # source_instance.soft_delete = True
                # destination_instance.soft_delete = True

                # Save changes
                manual_rate_instance.save()
                # version_rate_instance.save()
                # rate_instance.save()
                # company_instance.save()
                # source_instance.save()
                # destination_instance.save()

                return Response({'message': 'Records soft-deleted successfully'}, status=status.HTTP_204_NO_CONTENT)

        except ManualRate.DoesNotExist:
            return Response({"detail": "ManualRate not found."}, status=status.HTTP_404_NOT_FOUND)

        # except VersionedRate.DoesNotExist:
        #     return Response({"detail": "VersionedRate not found."}, status=status.HTTP_404_NOT_FOUND)

        # except Rate.DoesNotExist:
        #     return Response({"detail": "Rate not found."}, status=status.HTTP_404_NOT_FOUND)

        # except Company.DoesNotExist:
        #     return Response({"detail": "Company not found."}, status=status.HTTP_404_NOT_FOUND)

        # except Source.DoesNotExist:
        #     return Response({"detail": "Source not found."}, status=status.HTTP_404_NOT_FOUND)

        # except Destination.DoesNotExist:
        #     return Response({"detail": "Destination not found."}, status=status.HTTP_404_NOT_FOUND)

        except Exception as e:
            return Response({'detail': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        
class UpadatingRateFrozenInfoListView(APIView):
    permission_classes = [IsAuthenticated,IsClientUserEditAndRead|IsSystemOrClientAdmin|IsClientUserReadOnly]
    # permission_classes = [IsAuthenticated]

    def put(self, request, unique_uuid):
        requestData = request.data
        isRateUsed = requestData.get('isRateUsed', None)

        if isRateUsed is None:
            return Response({"error": "isRateUsed field is required"}, status=status.HTTP_400_BAD_REQUEST)

        # Fetch the ManualRate object (assuming unique_uuid refers to ManualRate's primary key)
        try:
            manual_rate = ManualRate.objects.get(unique_uuid=unique_uuid)
        except ManualRate.DoesNotExist:
            return Response({"error": "ManualRate not found"}, status=status.HTTP_404_NOT_FOUND)

        # Fetch the related VersionedRate and Rate objects based on the correct relationship field
        # try:
        #     versioned_rate = VersionedRate.objects.get(manual_rate=manual_rate)  # Adjust if needed
        #     rate = Rate.objects.get(manual_rate=manual_rate)  # Adjust if needed
        # except (VersionedRate.DoesNotExist, Rate.DoesNotExist):
        #     return Response({"error": "VersionedRate or Rate not found"}, status=status.HTTP_404_NOT_FOUND)

        # Update VersionedRate and Rate (common in both cases)
        # versioned_rate.isRateUsed = isRateUsed
        # versioned_rate.save()

        # rate.isRateUsed = isRateUsed
        # rate.save()

        # Update ManualRate
        manual_rate.isRateUsed = isRateUsed
        manual_rate.save()

        return Response({"message": "isRateUsed updated successfully"}, status=status.HTTP_200_OK)

class CustomerInfoListView(APIView):
    permission_classes = [IsAuthenticated,IsSystemOrClientAdmin|IsClientUserEditAndRead]
    # permission_classes = [IsAuthenticated]

    #  FOR GET 
    def get(self, request):
        try:
            customer_info = CustomerInfo.objects.all()
            customer_info_serializer = CustomerInfoSerializer(customer_info, many=True)
            return Response(customer_info_serializer.data, status=status.HTTP_200_OK)
        except CustomerInfo.DoesNotExist:
            return Response({"error": "Something went wrong"}, status=status.HTTP_404_NOT_FOUND)
        

    # FOR POST
    def post(self, request):
        try:
            requestData = request.data
            company_name = requestData.get('company_name')
            cust_name = requestData.get('cust_name')
            cust_email = requestData.get('cust_email')
            sales_represent = requestData.get('sales_represent')
            phone = requestData.get('phone')
            percentage = requestData.get('percentage')
            terms_condition = requestData.get('terms_condition')
            
            CustomerInfo.objects.create(
            company_name=company_name,
            cust_name=cust_name,
            cust_email=cust_email,
            sales_represent=sales_represent,
            phone=phone,
            percentage=percentage,
            terms_condition=terms_condition,
            )

            return Response({'message': 'Customer created successfully'}, status=status.HTTP_201_CREATED)

            
        except Exception as err:
            return Response({'details': str(err)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        
    
     #  FOR UPDATE BY ID 
    def put(self, request,id):
            try:
                requestData = request.data
                try:
                    customer_instance = CustomerInfo.objects.get(id=id)
                except CustomerInfo.DoesNotExist:
                    return Response({"detail": "Customer not found."}, status=status.HTTP_404_NOT_FOUND)

                customer_instance.percentage= requestData.get('percentage', customer_instance.percentage)
               
                 # Save changes
                customer_instance.save()
                # Return success response
                return Response({"detail": "Customer percentage updated successfully."}, status=status.HTTP_200_OK) 
            
            except CustomerInfo.DoesNotExist:
                return Response({"error": "Customer not found"}, status=status.HTTP_404_NOT_FOUND)  

class CutomerInfoDetailsListView(APIView):
     #  FOR GET BY ID 
    def get(self, request, id):
        try:
            # Retrieve the object based on the provided id
            customer_info = CustomerInfo.objects.get(id=id)
            
            # Serialize the object
            customer_info_serializer = CustomerInfoSerializer(customer_info)
            
            # Return the serialized data
            return Response(customer_info_serializer.data, status=status.HTTP_200_OK)
        except CustomerInfo.DoesNotExist:
            # Handle case where object with the given id does not exist
            return Response({"error": f"CustomerInfo with id {id} does not exist."}, status=status.HTTP_404_NOT_FOUND)
        except Exception as e:
            # Handle any other errors
            return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        
class RegistrationInfoListView(APIView):
    permission_classes = [IsAuthenticated,IsSystemOrClientAdmin]
    # permission_classes = [IsAuthenticated]

    #  FOR GET 
    def get(self, request):
        customer_info = Registration.objects.all()
        customer_info_serializer = RegistrationSerializer(customer_info, many=True)
        return Response(customer_info_serializer.data)
    
    # FOR POST
    def post(self, request):
        try:
            requestData = request.data
            name = requestData.get('name')
            email = requestData.get('email')
            username = requestData.get('username')
            phone = requestData.get('phone')
            password = requestData.get('password')
            
            
            Registration.objects.create(
            name=name,
            email=email,
            username=username,
            phone=phone,
            password=password
           
            )
            return Response({'message': 'Registration successfully'}, status=status.HTTP_201_CREATED)
            
        except Exception as err:
            return Response({'details': str(err)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
            
class CommodityList(generics.ListCreateAPIView):
    permission_classes=[IsAuthenticated,IsClientUserEditAndRead|IsSystemOrClientAdmin|IsClientUserReadOnly|IsUser]

    queryset = Comodity.objects.all()
    serializer_class = CommoditySerializer

class IncoTermList(generics.ListCreateAPIView):
    permission_classes=[IsAuthenticated,IsClientUserEditAndRead|IsSystemOrClientAdmin|IsClientUserReadOnly|IsUser]

    queryset = IncoTerm.objects.all()
    serializer_class = IncoTermSerializer

# ACTIVITY LOG FUNCTION
class ActivityLogView(APIView):
    # permission_classes = [IsAuthenticated|IsSystemAdministrator]
    permission_classes = [IsAuthenticated]

    # def get(self, request):
    #     try:
    #         activitityLogList = ActivityLog.objects.all()
    #         activitityLogListSerializer = ActivityLogSerializer(activitityLogList, many=True)
    #         return Response(activitityLogListSerializer.data)

    #     except Exception as err:
    #         return Response("Something went wrong")  


    def get(self, request):
        try:
            # Fetching filters from query parameters (if provided)
            recent_only = request.query_params.get('recent', 'false').lower() == 'true'
            user_id = request.query_params.get('user_id')  # Optional filter by user

            # If recent_only is enabled, fetch logs for the past 7 days (or modify this duration as needed)
            if recent_only:
                seven_days_ago = datetime.now() - timedelta(days=7)
                activitityLogList = ActivityLog.objects.filter(
                    Q(created_at__gte=seven_days_ago) &
                    (Q(userId=user_id) if user_id else Q())
                ).order_by('-created_at')
            else:
                activitityLogList = ActivityLog.objects.all().order_by('-created_at')

            activitityLogListSerializer = ActivityLogSerializer(activitityLogList, many=True)
            return Response(activitityLogListSerializer.data)

        except Exception as err:
            return Response({"error": "Something went wrong", "details": str(err)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


    def post(self, request):
        try:
            with transaction.atomic():
                requestData = request.data
                print(requestData)
                userId = requestData.get('userId')
                action_type = requestData.get('action_type')
                action_status = requestData.get('action_status')
                description = requestData.get('description')
                source_id = requestData.get('source_id')
                destination_id = requestData.get('destination_id')
                
                ActivityLog.objects.create(
                userId=userId,
                action_type=action_type,
                action_status=action_status,
                description=description,
                source_id=source_id,
                destination_id=destination_id,
                )
            return Response({'message': 'Log created successfully'}, status=status.HTTP_201_CREATED)    

        except Exception as err:
            return Response("Something went wrong")        