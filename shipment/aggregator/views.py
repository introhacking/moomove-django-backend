# views.py in your Django app
from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import status
import pdfplumber
import pandas as pd
import numpy as np
from .models import *
from .serializers import *
from pdfplumber import open as open_pdf
import PyPDF2
from django.http import JsonResponse
from django.views import View
from django.views.decorators.csrf import csrf_exempt
import re
# import fitz
from django.db import transaction
from rest_framework.parsers import MultiPartParser, FormParser, JSONParser
from docx import Document
from rest_framework import generics
from django.shortcuts import get_object_or_404
from django.contrib.auth import authenticate
from rest_framework_simplejwt.tokens import RefreshToken
from rest_framework.decorators import api_view
from rest_framework.permissions import IsAuthenticated
from drf_yasg import openapi

from rest_framework.permissions import AllowAny
from drf_spectacular.utils import extend_schema
from drf_yasg.utils import swagger_auto_schema
from google.oauth2 import service_account
from google.cloud import documentai_v1beta3 as documentai
from django.db import connection
from decimal import Decimal
import uuid
from dotenv import load_dotenv, dotenv_values
from rest_framework.authtoken.models import Token

from rest_framework.exceptions import ValidationError
from uauth.serializers import *
from uauth.models import *
from uauth.views import *
from uauth.role_permission import *
from datetime import datetime, timedelta, date
from django.utils.timezone import now
from django.db.models import Q
from django.db import transaction, IntegrityError
from django.core.exceptions import PermissionDenied
from django.core.exceptions import ObjectDoesNotExist
import json
import logging
logger = logging.getLogger(__name__)

config={
    **dotenv_values('constant_env/.env.shared'),
    **dotenv_values('constant_env/.env.secret'),
    **dotenv_values('constant_env/.env.error'),
}

# required_columns = [item.strip() for item in config.get("REQUIRED_COLUMN", '').split(",")]
# print(required_columns)
# print(config.get("REQUIRED_COLUMN" , "").split(","))
@extend_schema(
    request=UserSerializer,
    responses={200: 'application/json'},
    tags=["Authentication"],
    summary="Login to get JWT token",
    description="Provide username and password to receive JWT access and refresh tokens."
)

class ImportExcelData(APIView):
    permission_classes=[IsAuthenticated,IsSystemOrClientAdmin|IsClientUserEditAndRead]

    def post(self, request, format=None):
        file_obj = request.FILES.get('file')  # Assuming file is sent in the request
        if not file_obj:
            return Response({"error": config['ERROR_UPLOADING']}, status=status.HTTP_400_BAD_REQUEST)

        company_id = request.data.get('company_id')
        client_id = request.data.get('client_id')  # Extract client_id from the request
        # [ ADDED ON 29 / JAN/ 25 ]
        logger.info(f"Logged-in user: {request.user}, client_id: {client_id}, company_id :{company_id}")

        if not company_id:
            return Response({"error": config['ID_REQUIRED']}, status=status.HTTP_400_BAD_REQUEST)
        
        if not client_id:
            return Response({"error": "Client ID is required"}, status=status.HTTP_400_BAD_REQUEST)
        try:
            # company = Company.objects.get(id=company_id)  # by manish
            # Validate that the company belongs to the provided client_id
            company = ClientTemplateCompany.objects.filter(id=company_id, client_id=client_id).first()

            required_columns = ["Origin Port", "Destination Port", "Transit\ntime", "20'GP", "40'HC", "Effective Date", "Expiration Date"]
            sheets_to_read = ['F.E', 'E.Africa', 'Gulf-Red Sea']  # Adjust sheet names as per your Excel file

            # sheets_to_read = ['E.Africa']  # Adjust sheet names as per your Excel file

            # required_columns = config.get("REQUIRED_COLUMN" , "").split(",")
            # sheets_to_read = config.get("SHEETS_TO_READ" , "").split(",")  # Adjust sheet names as per your Excel file
            # sheets_to_read = ['E.Africa']  # Adjust sheet names as per your Excel file

            combined_df = pd.DataFrame()

            for sheet_name in sheets_to_read:
                df = pd.read_excel(file_obj, sheet_name=sheet_name, header=7, usecols=required_columns)
                df.dropna(how='all', inplace=True)
                df.rename(columns={"Transit\ntime": "Transit time"}, inplace=True)

                # Convert "Destination Port" to uppercase
                df['Destination Port'] = df['Destination Port'].str.upper().str.replace('PORT', '').str.strip()  # 4/Jan/2025
                # df['Destination Port'] = df['Destination Port'].str.replace('PORT', '').str.strip()

                # Convert Timestamps to string for JSON serialization
                df['Effective Date'] = df['Effective Date'].apply(lambda x: x.strftime('%Y-%m-%d') if isinstance(x, datetime) else None)
                df['Expiration Date'] = df['Expiration Date'].apply(lambda x: x.strftime('%Y-%m-%d') if isinstance(x, datetime) else None)

                # Append the data from the current sheet to the combined DataFrame
                combined_df = pd.concat([combined_df, df], ignore_index=True)

            # Replace NaN or NaT values with None for JSON serialization
            combined_df.replace({np.nan: None, pd.NaT: None}, inplace=True)

            results = combined_df.to_dict(orient='records')
            # Save the data to the database
            freight_types = ["20'GP", "40'HC"]

            for item in results:
                source_name = item["Origin Port"]
                destination_name = item["Destination Port"]
                transit_time_value = item["Transit time"]
                effective_date = item["Effective Date"]
                expiration_date = item["Expiration Date"]

                # source, _ = Source.objects.get_or_create(name=source_name)
                # destination, _ = Destination.objects.get_or_create(name=destination_name)
                # transit_time, _ = TransitTime.objects.get_or_create(time=str(transit_time_value))
                
                #[ 31/JAN/25]
                source, _ = Source.objects.get_or_create(name=source_name, client_id=company.client_id ) #client_id=client_id
                destination, _ = Destination.objects.get_or_create(name=destination_name, client_id=company.client_id  )
                # transit_time, _ = TransitTime.objects.get_or_create(time=str(transit_time_value ,client_id=company.client_id ))
                transit_time, _ = TransitTime.objects.get_or_create(time=str(transit_time_value), defaults={"client_id": company.client_id})

                for ft in freight_types:
                    rate_value = item.get(ft)

                    if rate_value is not None:
                        freight_type, _ = FreightType.objects.get_or_create(type=ft)

                        # Check if the rate record already exists
                        existing_rate = Rate.objects.filter(
                            company=company,
                            source=source,
                            destination=destination,
                            freight_type=freight_type,
                            client_id=client_id 
                        ).first()

                        # Fetch all versions sorted by effective_date descending
                        existing_versions = VersionedRate.objects.filter(
                            company=company,
                            source=source,
                            destination=destination,
                            freight_type=freight_type,
                            client_id=client_id 
                        ).order_by('-id')
                        for version in existing_versions:
                            print(f"Existing version found: {version}")
            
                        if existing_rate:
                            # Exclude the current existing_rate's version from existing_versions
                            existing_versions = existing_versions.exclude(id=existing_rate.version.id)

                            # Print existing versions for debugging
                            # for version in existing_versions:
                            #     print(f"Existing version found: {version}")

                            # Select the second latest version for mapping
                            second_latest_version = existing_versions.first() if existing_versions.exists() else None
                            # print("second_latest_version:", second_latest_version)

                            # Check if there are changes in rate value, effective date, or expiration date
                            has_changes = (
                                existing_rate.rate != rate_value or
                                existing_rate.effective_date != effective_date or
                                existing_rate.expiration_date != expiration_date
                            )

                            if has_changes:
                                # print("Inside has_changes")

                                # Create a new versioned rate
                                VersionedRate.objects.create(
                                    company=company,
                                    source=source,
                                    destination=destination,
                                    transit_time=transit_time,
                                    freight_type=freight_type,
                                    rate=rate_value,
                                    effective_date=effective_date,
                                    expiration_date=expiration_date,
                                    is_current=False,
                                    # [ 31/JAN/25]
                                    client_id=client_id
                                )

                                # Update the existing rate
                                existing_rate.rate = rate_value
                                existing_rate.effective_date = effective_date
                                existing_rate.expiration_date = expiration_date
                                if second_latest_version:
                                    existing_rate.version = second_latest_version

                                existing_rate.save()
                        else:
                            # print("inside else")

                            # Create new rate and version
                            versioned_rate = VersionedRate.objects.create(
                                company=company,
                                source=source,
                                destination=destination,
                                transit_time=transit_time,
                                freight_type=freight_type,
                                rate=rate_value,
                                effective_date=effective_date,
                                expiration_date=expiration_date,
                                is_current=True,
                                # [ 31/JAN/25]
                                client_id=client_id
                            )
                            Rate.objects.create(
                                company=company,
                                source=source,
                                destination=destination,
                                transit_time=transit_time,
                                freight_type=freight_type,
                                rate=rate_value,
                                effective_date=effective_date,
                                expiration_date=expiration_date,
                                version=versioned_rate,
                                # [ 31/JAN/25]
                                client_id=client_id
                            )
            return Response({"message": f"Excel {config['SUCCESS_UPLOADED_MESSAGE']}", "results": results}, status=status.HTTP_201_CREATED)

        except ClientTemplateCompany.DoesNotExist:
            return Response({"error": "Company does not exist"}, status=status.HTTP_400_BAD_REQUEST)
        except Exception as e:
            return Response({"error": str(e)}, status=status.HTTP_400_BAD_REQUEST)
    
class ExtractWordTableView(APIView):
    parser_classes = [MultiPartParser]
    permission_classes=[IsAuthenticated,IsSystemOrClientAdmin|IsClientUserEditAndRead]
    # permission_classes=[IsAuthenticated]

    def post(self, request, *args, **kwargs):
        try:
            # Assume the file is passed as 'file' in the POST request
            file_obj = request.FILES['file']

            # Extract company ID from the request data
            company_id = request.data.get('company_id')
            client_id = request.data.get('client_id')
            
            logger.info(f"Logged-in user: {request.user}, client_id: {client_id}, company_id :{company_id}")
            if not client_id:
                return JsonResponse({"error": "Client ID is required"}, status=status.HTTP_400_BAD_REQUEST)
            
            # Extract company ID from the request data
            if not company_id:
                return JsonResponse({"error": config['ID_REQUIRED']}, status=status.HTTP_400_BAD_REQUEST)

            # Process the uploaded Word document
            extracted_data, expiration_date = self.extract_table_data(file_obj)

            # Convert extracted data to desired format
            converted_data = [self.convert_to_desired_format(item, expiration_date) for item in extracted_data]

            # Save the converted data
            self.save_imported_data(converted_data, company_id)

            # Return JSON response
            return Response({"message": f"Word {config['SUCCESS_UPLOADED_MESSAGE']}", "results": converted_data}, status=status.HTTP_201_CREATED)

        except Exception as e:
            # Handle exceptions (e.g., file not found, invalid format, etc.)
            return JsonResponse({"error": str(e)}, status=status.HTTP_400_BAD_REQUEST)

    def extract_table_data(self, file_obj):
        doc = Document(file_obj)
        tables = doc.tables

        # Assuming the table you want is the first table in the document
        table = tables[0]

        data = []
        keys = None

        for i, row in enumerate(table.rows):
            text = [cell.text.strip() for cell in row.cells]
            if i == 0:
                keys = text  # assuming first row as header
            else:
                data.append(dict(zip(keys, text)))
        
        # Extracting the "Valid till" date from the document
        expiration_date = self.extract_valid_till_date(doc)
        
        return data, expiration_date

    def extract_valid_till_date(self, doc):
        valid_till_pattern = re.compile(r"Valid till\s+(\d{2}/\d{2}\s*/\s*\d{4})", re.IGNORECASE)
        
        for paragraph in doc.paragraphs:
            match = valid_till_pattern.search(paragraph.text)
            if match:
                    date_str = match.group(1).replace(" ", "")  # Remove any spaces within the date
                    # Convert the extracted date to a datetime object
                    date_obj = datetime.strptime(date_str, "%d/%m/%Y")
                    # Format the date to YYYY-MM-DD
                    return date_obj.strftime("%Y-%m-%d")  # Remove any spaces within the date
        
        return ""

    def convert_to_desired_format(self, data, expiration_date):
        formatted_data = {
            "Origin Port": "Nhava Sheva",
            "Destination Port": data.get("Sector", ""),
            "Transit time": data.get("Appx.T.time", 0),
            "20'GP": float(data.get("20'GP", 0)),
            "40'HC": float(data.get("40'GP/HC", 0)),
            "Effective Date": "2024-04-01",
            "Expiration Date": expiration_date if expiration_date else data.get("Expiration Date", "")
        }
        return formatted_data

    def save_imported_data(self, results, company_id):
        client_id = self.request.data.get('client_id')  # Assuming client_id is passed in the request

        if not client_id:
            raise ValueError("Client ID is required for filtering rates.")
        freight_types = ["20'GP", "40'HC"]  # Define the freight types you are processing
        company = get_object_or_404(ClientTemplateCompany, id=company_id)
        for item in results:
            source_name = item["Origin Port"]
            destination_name = item["Destination Port"]
            transit_time_value = item["Transit time"]
            effective_date = item["Effective Date"]
            expiration_date = item["Expiration Date"]
            print(f"Source: {source_name}, Destination: {destination_name}, Transit Time: {transit_time_value}, "
                f"Effective Date: {effective_date}, Expiration Date: {expiration_date}")
            source, _ = Source.objects.get_or_create(name=source_name, client_id=company.client_id)
            destination, _ = Destination.objects.get_or_create(name=destination_name, client_id=company.client_id)
            transit_time, _ = TransitTime.objects.get_or_create(time=str(transit_time_value), defaults={"client_id": company.client_id})
    
            for ft in freight_types:
                rate_value = item.get(ft)
                if rate_value is not None:
                    # company = get_object_or_404(Company, id=company_id) # by manish
                    # company = get_object_or_404(ClientTemplateCompany, id=company_id)

                    freight_type, _ = FreightType.objects.get_or_create(type=ft)

                    # Check if the rate record already exists
                    existing_rate = Rate.objects.filter(
                        company=company,
                        source=source,
                        destination=destination,
                        freight_type=freight_type,
                        client_id=client_id
                    ).first()

                    # Fetch all versions sorted by effective_date descending
                    existing_versions = VersionedRate.objects.filter(
                        company=company,
                        source=source,
                        destination=destination,
                        freight_type=freight_type,
                        client_id=client_id 

                    ).order_by('-id')
                    # Exclude the current existing_rate
                    if existing_rate:
                        existing_versions = existing_versions.exclude(id=existing_rate.version.id)
                        for version in existing_versions:
                            print(f"Existing version found: {version}")

                    # Select the second latest version for mapping
                    second_latest_version = existing_versions[0] if len(existing_versions) > 0 else None
                    print("second_latest_version,second_latest_version",  second_latest_version)

                    if existing_rate:
                        has_changes = (
                            existing_rate.rate != rate_value or
                            existing_rate.effective_date != effective_date or
                            existing_rate.expiration_date != expiration_date
                        )

                        if has_changes:
                            VersionedRate.objects.create(
                                company=company,
                                source=source,
                                destination=destination,
                                transit_time=transit_time,
                                freight_type=freight_type,
                                rate=rate_value,
                                effective_date=effective_date,
                                expiration_date=expiration_date,
                                is_current=False,
                                client_id=client_id
                            )
                            # Update the existing rate
                            existing_rate.rate = rate_value
                            existing_rate.effective_date = effective_date
                            existing_rate.expiration_date = expiration_date
                            if second_latest_version:
                                existing_rate.version = second_latest_version

                            existing_rate.save()
                    else:
                        # Create new rate and version
                        versioned_rate = VersionedRate.objects.create(
                            company=company,
                            source=source,
                            destination=destination,
                            transit_time=transit_time,
                            freight_type=freight_type,
                            rate=rate_value,
                            effective_date=effective_date,
                            expiration_date=expiration_date,
                            is_current=True,
                            client_id=client_id
                        )
                        Rate.objects.create(
                            company=company,
                            source=source,
                            destination=destination,
                            transit_time=transit_time,
                            freight_type=freight_type,
                            rate=rate_value,
                            effective_date=effective_date,
                            expiration_date=expiration_date,
                            version=versioned_rate,
                            client_id=client_id
        )

class ExtractPDFTableView(APIView):
    permission_classes=[IsAuthenticated,IsSystemOrClientAdmin|IsClientUserEditAndRead]
    # permission_classes=[IsAuthenticated]

 
    def extract_valid_dates(self, text):
        valid_dates = {}

        # Regex pattern to find valid dates and regions
        pattern = r"valid(?:\s+FROM)?\s+(\d{2}-\d{2}-\d{4})\s+to\s+(\d{2}-\d{2}-\d{4})\s+(Vessel Sailing\s+)?([A-Z\s-]+)"
        matches = re.finditer(pattern, text, re.IGNORECASE)

        for match in matches:
            start_date = datetime.strptime(match.group(1), "%d-%m-%Y").date()
            end_date = datetime.strptime(match.group(2), "%d-%m-%Y").date()
            region = match.group(4).strip()

            if '\n' in region:
                region = region.split('\n', 1)[0].strip()

            # Remove "Vessel Sailing" prefix if present
            if region.startswith("Vessel Sailing"):
                region = region[len("Vessel Sailing"):].strip()

            valid_dates[region] = {"valid_from": start_date, "valid_to": end_date}


        return valid_dates
    
    def process_document(self, content, mime_type):
        KEYFILE_PATH = r'E:\PROJECT\MooMove\app\shipment\aggregator\jspl-trocr-2024-419309-21a5c31d8f62.json'
        # Replace these with your actual Google Cloud project ID and processor ID
        project_id = config['PROJECT_ID']
        location = config['LOCATION']
        processor_id = config['PROCESSOR_ID']

        # Authenticate with service account JSON key file
        credentials = service_account.Credentials.from_service_account_file(KEYFILE_PATH)
        client = documentai.DocumentProcessorServiceClient(credentials=credentials)

        # The full resource name of the processor
        name = f'projects/{project_id}/locations/{location}/processors/{processor_id}'

        # Configure the process request
        request = {"name": name, "raw_document": {"content": content, "mime_type": mime_type}}

        # Recognize text entities in the document
        result = client.process_document(request=request)

        document = result.document

        # Get the document text
        document_text = document.text
        # print("document_text: ",document_text)
        # Get individual entities
        entities_text = ""
        for entity in document.entities:
            entities_text += f"<li>Entity type: {entity.type_}, Value: {entity.mention_text}</li>"

        return document_text, entities_text
    
    def clean_region_name(self, region_name):
        # Clean up region name to match keys in valid_dates
        return region_name.replace('\n', '').strip()


    def apply_valid_dates(self, tables, valid_dates):
        for table in tables:
            region = self.clean_region_name(table.get('Region', ''))
            

            if region in valid_dates:
                table['Valid From'] = valid_dates[region]['valid_from']
                table['Valid To'] = valid_dates[region]['valid_to']
            elif region  == "WEST AFRICA":
                # Assign same dates as SOUTH WEST AFRICA
                south_west_africa_dates = valid_dates.get('SOUTH WEST AFRICA')
                if south_west_africa_dates:
                    table['Valid From'] = south_west_africa_dates['valid_from']
                    table['Valid To'] = south_west_africa_dates['valid_to']
            else:
                table['Valid From'] = None
                table['Valid To'] = None

    def post(self, request):

        try:
            # Check if the 'file' key is present in request.FILES
            if 'file' not in request.FILES:
                return Response({"error": "No 'file' key found in request data"}, status=status.HTTP_400_BAD_REQUEST)

            # Extract company_id from request data or session if authenticated
            company_id = request.data.get('company_id')  # Adjust this based on how company_id is passed
            client_id=request.data.get('client_id')# Adjust this based on how company_id is passed
            
            logger.info(f"Logged-in user: {request.user}, client_id: {client_id}")

            # Retrieve company instance based on company_id
            # company = Company.objects.get(id=company_id)  # Assuming Company model and id field exist   # by manish
            company = ClientTemplateCompany.objects.get(id=company_id,client_id=client_id)  # Assuming Company model and id field exist

            pdf_file = request.FILES['file']
            content = pdf_file.read()
            mime_type = pdf_file.content_type

            # Process the document
            document_text, entities_text = self.process_document(content, mime_type)
            all_tables = []
            regions = {
                "WEST AFRICA": ["LOME", "TEMA", "ABIDJAN", "COTONOU", "TINCAN ISLAND", "LAGOS", "ONNE", "DAKAR"],
                "SOUTH WEST AFRICA": ["LUANDA", "POINTE NOIRE", "DOUALA", "MATADI", "WALVIS BAY", "KRIBI"],
                "EAST AFRICA": ["MOMBASA", "DAR-ES-SALAAM"],
                "SOUTH AFRICA": ["DURBAN", "CAPETOWN"]
            }

            reader = PyPDF2.PdfReader(pdf_file)
            num_pages = len(reader.pages)

            for page_num in range(num_pages):
                page = reader.pages[page_num]
                text = page.extract_text()
                # print("dockumet text: ", document_text)
                # print("text: ", text)
                # Extract valid dates for each region
                valid_dates = self.extract_valid_dates(document_text)
                

                # Regex pattern for main table rows like 'MATADI 2075 3450 2075 3450'
                table_pattern = r"([A-Z\s-]+)\s+(\d+)\s+(\d+)\s+(\d+)\s+(\d+)"
                table_matches = re.finditer(table_pattern, text, re.MULTILINE)

                # Iterate through each match and capture row number
                for idx, match in enumerate(table_matches, start=1):
                    row_number = match.start()
                    # port_name = match.group(1).strip().replace("PORT", "").strip()  # Remove "PORT"
                    # port_name = match.group(1).split('\n', 1)[-1].strip().replace("PORT", "").strip()
                    port_name = match.group(1).split('\n', 1)[-1].strip().replace("PORT", "").strip().replace("  ", " ").replace(" -", "-")


                    rates_ex_nhava_sheva = {
                        "origin_port": "Nhava Sheva",
                        "20'GP": int(match.group(2)),
                        "40'HC": int(match.group(3))
                    }
                    rates_ex_mundra = {
                        "origin_port": "Mundra",
                        "20'GP": int(match.group(4)),
                        "40'HC": int(match.group(5))
                    }

                    # Determine region based on destination port
                    region = None
                    for key, value in regions.items():
                        if any(port in port_name for port in value):
                            region = key
                            break

                    table_data = {
                        "Row Number": row_number,
                        "Destination Port": port_name,
                        "Origin Port Nhava Sheva": rates_ex_nhava_sheva,
                        "Origin Port Mundra": rates_ex_mundra,
                        "Transit time": None,  # Initialize as None
                        "Region": region,  # Assign region
                        "Valid From": valid_dates[region]["valid_from"] if region in valid_dates else None,
                        "Valid To": valid_dates[region]["valid_to"] if region in valid_dates else None
                    }
                     
                    all_tables.append(table_data)
                self.apply_valid_dates(all_tables, valid_dates)
                # Regex pattern to find transit time like '42 days'
                days_pattern = r"(\d+)\s+days"
                days_matches = re.finditer(days_pattern, text)

                # Find the corresponding transit time and update in all_tables
                for days_match in days_matches:
                    transit_time = int(days_match.group(1))

                    # Find the nearest row number before the transit time match
                    row_number = -1
                    for table_data in all_tables:
                        if table_data["Row Number"] < days_match.start():
                            row_number = table_data["Row Number"]
                        else:
                            break

                    # Update transit time in the corresponding row
                    if row_number != -1:
                        for table_data in all_tables:
                            if table_data["Row Number"] == row_number:
                                table_data["Transit time"] = transit_time
                                break
            if all_tables:
                created_records = []

                for item in all_tables:
                    destination_port = item['Destination Port']
                    transit_time = item['Transit time']
                    region = item['Region']
                    valid_from = item['Valid From']
                    valid_to = item['Valid To']

                    origin_ports = [item['Origin Port Nhava Sheva'], item['Origin Port Mundra']]

                    for origin_port in origin_ports:
                        record = {
                            "Destination Port": destination_port,
                            "Origin Port": origin_port['origin_port'] if 'origin_port' in origin_port else origin_port['origin_port'],
                            "20'GP": origin_port["20'GP"],
                            "40'HC": origin_port["40'HC"],
                            "Transit time": transit_time,
                            "effective_date": valid_from,
                            "expiration_date": valid_to
                        }

                        created_records.append(record)

                # Save created_records to database
                for item in created_records:
                    source_name = item["Origin Port"]
                    destination_name = item["Destination Port"]
                    transit_time_value = item["Transit time"]
                    effective_date = item["effective_date"]
                    expiration_date = item["expiration_date"]

                    source, _ = Source.objects.get_or_create(name=source_name)
                    destination, _ = Destination.objects.get_or_create(name=destination_name)
                    transit_time, _ = TransitTime.objects.get_or_create(time=str(transit_time_value))

                    freight_types = ["20'GP", "40'HC"]
                    for ft in freight_types:
                        rate_value = item.get(ft)

                        if rate_value is not None:
                            freight_type, _ = FreightType.objects.get_or_create(type=ft)

                            # Check if the rate record already exists
                            existing_rate = Rate.objects.filter(
                                company=company,
                                source=source,
                                destination=destination,
                                freight_type=freight_type,
                                client_id=client_id 
                            ).first()

                            # Fetch all versions sorted by effective_date descending
                            existing_versions = VersionedRate.objects.filter(
                                company=company,
                                source=source,
                                destination=destination,
                                freight_type=freight_type,
                                client_id=client_id 
                            ).order_by('-id')
                            for version in existing_versions:
                                print(f"Existing version found1: {version}")

                            if existing_rate:
                                # Exclude the current existing_rate's version from existing_versions
                                existing_versions = existing_versions.exclude(id=existing_rate.version.id)

                                # Print existing versions for debugging
                                for version in existing_versions:
                                    print(f"Existing version found: {version}")

                                # Select the second latest version for mapping
                                second_latest_version = existing_versions.first() if existing_versions.exists() else None
                                print("second_latest_version:", second_latest_version)

                                # Check if there are changes in rate value, effective date, or expiration date
                                has_changes = (
                                    existing_rate.rate != rate_value or
                                    existing_rate.effective_date != effective_date or
                                    existing_rate.expiration_date != expiration_date
                                )

                                if has_changes:
                                    print("Inside has_changes")

                                    # Create a new versioned rate
                                    VersionedRate.objects.create(
                                        company=company,
                                        source=source,
                                        destination=destination,
                                        transit_time=transit_time,
                                        freight_type=freight_type,
                                        rate=rate_value,
                                        effective_date=effective_date,
                                        expiration_date=expiration_date,
                                        is_current=False
                                    )

                                    # Update the existing rate
                                    existing_rate.rate = rate_value
                                    existing_rate.effective_date = effective_date
                                    existing_rate.expiration_date = expiration_date
                                    if second_latest_version:
                                        existing_rate.version = second_latest_version

                                    existing_rate.save()
                            else:
                                # Create new rate and version
                                versioned_rate = VersionedRate.objects.create(
                                    company=company,
                                    source=source,
                                    destination=destination,
                                    transit_time=transit_time,
                                    freight_type=freight_type,
                                    rate=rate_value,
                                    effective_date=effective_date,
                                    expiration_date=expiration_date,
                                    is_current=True
                                )
                                Rate.objects.create(
                                    company=company,
                                    source=source,
                                    destination=destination,
                                    transit_time=transit_time,
                                    freight_type=freight_type,
                                    rate=rate_value,
                                    effective_date=effective_date,
                                    expiration_date=expiration_date,
                                    version=versioned_rate
                                )

                return Response({"message": f"PDF {config['SUCCESS_UPLOADED_MESSAGE']}", "tables": all_tables}, status=status.HTTP_200_OK)

        except KeyError:
            return Response({"error": "No 'file' key found in request data"}, status=status.HTTP_400_BAD_REQUEST)

        except Exception as e:
            return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

class RateWithVersionsAPIView(APIView):
    permission_classes = [
        IsAuthenticated,
        IsClientUserEditAndRead | IsSystemOrClientAdmin | IsClientUserReadOnly | IsUser | IsSuperAdmin,
    ]
    
    # [11/03/2025]
    def get(self, request, company_id):
        try:
            user = request.user
            # Determine the active client:
            # If the user is an admin and has switched (current_client is set), then use that;
            # Otherwise, use the user's assigned client.
            if user.is_admin and user.current_client:
                active_client_id = user.current_client.client_id
            else:
                # For non-admin users, we assume client is always set.
                if not user.client:
                    return Response(
                        {"error": "You are not associated with any client."},
                        status=status.HTTP_403_FORBIDDEN,
                    )
                active_client_id = user.client.client_id

            # Admin can access all rates for the selected company, regardless of client.
            # Here, you might choose to have a different logic if admin is allowed to see data for any client.
            if user.is_admin:
                rates = Rate.objects.filter(company_id=company_id, soft_delete=False)
            else:
                # For regular users, filter rates by the active client.
                rates = Rate.objects.filter(
                    company_id=company_id,
                    company__client_id=active_client_id,
                    soft_delete=False,
                )

            rate_serializer = RateSerializer(rates, many=True)
            return Response(rate_serializer.data, status=status.HTTP_200_OK)

        except Exception as e:
            return Response(
                {"error": f"An unexpected error occurred: {str(e)}"},
                status=status.HTTP_400_BAD_REQUEST,
            ) 




    # [ 18/feb/25]
    # def get(self, request, company_id):
    #     try:
    #         user = request.user

    #     # Super Admin can access all rates
    #         if user.is_admin or (user.role and user.role.role_name == "Super Admin"):
    #             rates = Rate.objects.filter(company_id=company_id, soft_delete=False)
    #         else:
    #         # Regular users can only access rates related to their client_id
    #             if not user.client_id:
    #                 return Response(
    #                 {"error": "You are not associated with any client."},
    #                 status=status.HTTP_403_FORBIDDEN,
    #             )

    #         rates = Rate.objects.filter(
    #             company_id=company_id,
    #             company__client_id=user.client_id,
    #             soft_delete=False,
    #         )

    #     # Serialize the rates
    #         rate_serializer = RateSerializer(rates, many=True)
    #         return Response(rate_serializer.data, status=status.HTTP_200_OK)

    #     except Exception as e:
    #         return Response(
    #         {"error": f"An unexpected error occurred: {str(e)}"},
    #         status=status.HTTP_400_BAD_REQUEST,
    #     )



class CompanyListAPIView(APIView):
    permission_classes=[IsAuthenticated,IsClientUserEditAndRead|IsSystemOrClientAdmin|IsClientUserReadOnly|IsUser|IsSuperAdmin]

    #[11/03/2025]
    def get(self, request):
        user = request.user

    # Determine the active client for filtering
        if user.is_admin and user.current_client:
            active_client_id = user.current_client.client_id
        else:
        # For non-admin users, use their assigned client
            active_client_id = user.client_id

    # Super Admin can access all companies regardless of client switching if desired,
    # but if you want them to be limited by the switched client, use active_client_id.
        if user.is_admin or (user.role and user.role.role_name == "Super Admin"):
        # If admin, you could either:
        # Option A: Allow access to all companies (current implementation)
        # companies = Company.objects.filter(soft_delete=False)
        #
        # Option B: Limit to the switched client (if current_client is set)
            companies = Company.objects.filter(
                client_id=active_client_id,
                soft_delete=False
                )
        else:
    # Regular users can only access companies related to their own client
            companies = Company.objects.filter(
            client_id=active_client_id,
            soft_delete=False
            )

        serializer = CompanySerializer(companies, many=True)
        return Response(serializer.data)



    #[18/feb/25]
    # def get(self, request):
    #     # Retrieve the user's client
    #     user = request.user

    # # Super Admin can access all companies
    #     if user.is_admin or (user.role and user.role.role_name == "Super Admin"):
    #         companies = Company.objects.filter(soft_delete=False)
    #     else:
    #     # Regular users can only access companies related to their client_id
    #         companies = Company.objects.filter(client_id=user.client_id, soft_delete=False)

    #     serializer = CompanySerializer(companies, many=True)
    #     return Response(serializer.data)


    # def post(self, request):
    #     serializer = CompanySerializer(data=request.data)
    #     if serializer.is_valid():
    #         serializer.save()
    #         return Response(serializer.data, status=status.HTTP_201_CREATED)
    #     return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


    # def get(self, request):
    #     # Retrieve the user's client
    #     client = request.user.client

    #     # Filter companies based on the user's client_id
    #     if client:
    #         companies = Company.objects.filter(client_id=client.client_id, soft_delete=False)
    #     else:
    #         # If the user does not belong to a specific client, retrieve all companies (for admins)
    #         companies = Company.objects.filter(soft_delete=False)

    #     serializer = CompanySerializer(companies, many=True)
    #     return Response(serializer.data)


class ClientTemplateCompanyAPIView(APIView):
    permission_classes = [
        IsAuthenticated,
        IsClientUserEditAndRead | IsSystemOrClientAdmin | IsClientUserReadOnly | IsUser | IsSuperAdmin,
    ]

    # [ 11/03/2025 ]
    def get(self, request):
        try:
            user = request.user

            # Determine the active client for filtering.
            # For admin users who have switched, use current_client.
            if user.is_admin and user.current_client:
                active_client_id = user.current_client.client_id
            else:
                # For regular users, ensure they have an associated client.
                if not user.client_id:
                    return Response(
                        {"error": "You are not associated with any client."},
                        status=status.HTTP_403_FORBIDDEN,
                    )
                active_client_id = user.client_id

            # For admin (or super admin), you may choose one of two approaches:
            # Option A: Allow them to access all companies regardless of client filtering.
            # Option B: Limit them to companies of the active client.
            # Here, Option B is applied.
            if user.is_admin or (user.role and user.role.role_name == "Super Admin"):
                companies = ClientTemplateCompany.objects.filter(
                    client_id=active_client_id, soft_delete=False
                )
            else:
                companies = ClientTemplateCompany.objects.filter(
                    client_id=active_client_id, soft_delete=False
                )

            serializer = ClientTemplateCompanySerializer(companies, many=True)
            return Response(serializer.data, status=status.HTTP_200_OK)

        except Exception as e:
            return Response(
                {"error": f"An unexpected error occurred: {str(e)}"},
                status=status.HTTP_400_BAD_REQUEST,
            )

    # [18/feb/25]
    # def get(self, request):
    #     try:
    #         user = request.user

    #     # Super Admin can access all companies
    #         if user.is_admin or (user.role and user.role.role_name == "Super Admin"):
    #             companies = ClientTemplateCompany.objects.filter(soft_delete=False)
    #         else:
    #         # Regular users can only access companies related to their client_id
    #             if not user.client_id:
    #                 return Response(
    #                 {"error": "You are not associated with any client."},
    #                 status=status.HTTP_403_FORBIDDEN,
    #             )

    #             companies = ClientTemplateCompany.objects.filter(
    #             client_id=user.client_id, soft_delete=False
    #         )

    #         serializer = ClientTemplateCompanySerializer(companies, many=True)
    #         return Response(serializer.data, status=status.HTTP_200_OK)

    #     except Exception as e:
    #         return Response(
    #         {"error": f"An unexpected error occurred: {str(e)}"},
    #         status=status.HTTP_400_BAD_REQUEST,
    #     )

    def post(self, request):
        try:
            # Check if the user is associated with a client
            user_client = request.user.client
            logger.info(f"Logged-in user: {request.user}, client_id:{user_client}")
            if not user_client:
                raise PermissionDenied("You are not associated with any client.")

            # Attach the client's ID to the request data
            request.data["client_id"] = user_client.client_id


            # Serialize and save the data
            serializer = ClientTemplateCompanySerializer(data=request.data)
            if serializer.is_valid():
                serializer.save()
                return Response(serializer.data, status=status.HTTP_201_CREATED)

            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

        except PermissionDenied as e:
            return Response({"error": str(e)}, status=status.HTTP_403_FORBIDDEN)
        except Exception as e:
            return Response(
                {"error": f"An unexpected error occurred: {str(e)}"},
                status=status.HTTP_400_BAD_REQUEST,
            )

class SourceListAPIView(APIView):
    permission_classes = [
        IsAuthenticated,
        IsClientUserEditAndRead | IsSystemOrClientAdmin | IsClientUserReadOnly | IsUser | IsSuperAdmin,
    ]

    # [11/03/2025]
    def get(self, request):
        user = request.user

    # Determine the active client for filtering.
        if user.is_admin and user.current_client:
            active_client_id = user.current_client.client_id
        else:
            active_client_id = user.client_id

    # For admin/super admin, you may choose to restrict to the active client if desired.
        if user.is_admin or (user.role and user.role.role_name == "Super Admin"):
            sources = Source.objects.filter(client_id=active_client_id, soft_delete=False)
        else:
            sources = Source.objects.filter(soft_delete=False, client_id=active_client_id)

        serializer = SourceSerializer(sources, many=True)
        return Response(serializer.data)
    


    # [18/feb/25]
    # def get(self, request):
    #     user = request.user

    #     # Super Admin can access all sources
    #     if user.is_admin or (user.role and user.role.role_name == "Super Admin"):
    #         sources = Source.objects.filter(soft_delete=False)
    #     else:
    #         sources = Source.objects.filter(soft_delete=False, client_id=user.client_id)

    #     serializer = SourceSerializer(sources, many=True)
    #     return Response(serializer.data)

     

    #  [ PDF POST ]
    # def post(self, request, *args, **kwargs):
    #     try:
    #         file = request.FILES.get('file')
    #         client_id = request.data.get('client_id')

    #         if not file:
    #             return Response({"error": "No file uploaded"}, status=status.HTTP_400_BAD_REQUEST)

    #         if not client_id:
    #             return Response({"error": "client_id is required"}, status=status.HTTP_400_BAD_REQUEST)

    #         # Save file temporarily
    #         file_path = default_storage.save(f"temp/{file.name}", file)

    #         # Process PDF
    #         with pdfplumber.open(default_storage.path(file_path)) as pdf:
    #             data_list = []
    #             for page in pdf.pages:
    #                 table = page.extract_table()
    #                 if table:
    #                     for row in table[1:]:  # Skip header
    #                         locode, country_code, country_name, port_code, port_name = row

    #                         if port_name and country_name:
    #                             combined_name = f"{port_name}, {country_name}"
    #                             data_list.append(Source(name=combined_name, client_id=client_id))

    #             # Bulk insert data
    #             Source.objects.bulk_create(data_list, ignore_conflicts=True)

    #         return Response({"success": f"{len(data_list)} records inserted"}, status=status.HTTP_201_CREATED)

    #     except Exception as e:
    #         return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR) 


class DestinationListAPIView(APIView):
    permission_classes = [
        IsAuthenticated,
        IsClientUserEditAndRead | IsSystemOrClientAdmin | IsClientUserReadOnly | IsUser | IsSuperAdmin,
    ]

    # [11/03/2025]
    def get(self, request):
        user = request.user

        # Determine the active client for filtering.
        if user.is_admin and user.current_client:
            active_client_id = user.current_client.client_id
        else:
            active_client_id = user.client_id

        if user.is_admin or (user.role and user.role.role_name == "Super Admin"):
            destination = Destination.objects.filter(client_id=active_client_id, soft_delete=False)
        else:
            destination = Destination.objects.filter(soft_delete=False, client_id=active_client_id)

        serializer = DestinationSerializer(destination, many=True)
        return Response(serializer.data)


    # [18/feb/25]
    # def get(self, request):
    #     user = request.user

    #     # Super Admin can access all sources
    #     if user.is_admin or (user.role and user.role.role_name == "Super Admin"):
    #         destination = Destination.objects.filter(soft_delete=False)
    #     else:
    #         destination = Destination.objects.filter(soft_delete=False, client_id=user.client_id)

    #     serializer = DestinationSerializer(destination, many=True)
    #     return Response(serializer.data)

   

class FreightTypeListAPIView(APIView):
    permission_classes = [
        IsAuthenticated,
        IsClientUserEditAndRead | IsSystemOrClientAdmin | IsClientUserReadOnly | IsUser | IsSuperAdmin,
    ]

    # [11/03/2025]
    def get(self, request):
        user = request.user

        # Determine the active client for filtering.
        if user.is_admin and user.current_client:
            active_client_id = user.current_client.client_id
        else:
            active_client_id = user.client_id

        if user.is_admin or (user.role and user.role.role_name == "Super Admin"):
            freightType = FreightType.objects.filter(client_id=active_client_id, soft_delete=False)
        else:
            freightType = FreightType.objects.filter(soft_delete=False, client_id=active_client_id)

        serializer = FreightTypeSerializer(freightType, many=True)
        return Response(serializer.data)
    
    
    # [18/feb/25]
    # def get(self, request):
    #     user = request.user

    #     # Super Admin can access all sources
    #     if user.is_admin or (user.role and user.role.role_name == "Super Admin"):
    #         freightType =FreightType.objects.filter(soft_delete=False)
    #     else:
    #         freightType = FreightType.objects.filter(soft_delete=False, client_id=user.client_id)

    #     serializer = FreightTypeSerializer(freightType, many=True)
    #     return Response(serializer.data)

  

    def post(self, request):
        try:
            # Fetch the client_id from the logged-in user
            client_id = request.user.client_id

            logger.info(f"Logged-in user: {request.user}, client_id:{client_id}")

            if not client_id:
                raise PermissionDenied("You are not associated with any client.")

            # Extract freight type from the request data
            request_data = request.data
            freight_type = request_data.get('type')

            # Check if a FreightType with this type already exists for the client
            existing_freight_type = FreightType.objects.filter(
                type=freight_type, client_id=client_id
            ).first()

            if existing_freight_type:
                return Response(
                    {"message": "Freight type already exists."}, 
                    status=status.HTTP_200_OK
                )

            # Create a new FreightType instance for the client
            FreightType.objects.create(type=freight_type, client_id=client_id)

            return Response(
                {'message': 'Freight type created successfully.'}, 
                status=status.HTTP_201_CREATED
            )

        except PermissionDenied as e:
            return Response({"error": str(e)}, status=status.HTTP_403_FORBIDDEN)
        except Exception as e:
            return Response(
                {'detail': f"An unexpected error occurred: {str(e)}"}, 
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )

class RateListView(generics.ListAPIView):
    permission_classes = [IsAuthenticated, IsClientUserEditAndRead | IsSystemOrClientAdmin | IsClientUserReadOnly | IsUser | IsSuperAdmin,]
    serializer_class = RateSerializer1 # [ NOTEE : ALL FIELDS ARE MENTIONS ALSO]
  
    # [11/MAR/2025]
    def get_queryset(self):
        source_id = self.kwargs.get('source')
        destination_id = self.kwargs.get('destination')
        freight_type_id = self.kwargs.get('freight_type', None)
        # [18/feb/25]
        user = self.request.user

        # Get the client_id of the authenticated user
        # client_id = self.request.user.client.client_id if self.request.user.client else None

        # if not client_id:
        #     raise PermissionDenied({"error": "You are not associated with any client."})

        # [18/feb/25]
        if user.is_admin or (user.role and user.role.role_name == "Super Admin"):
            if user.current_client:
                client_id = user.current_client.client_id
            else:
                # No active client switch; pass None to fetch all data for admin.
                client_id = None
              # Pass NULL in the stored procedure to fetch all data

        else: 
            client_id = user.client_id

        # Check if the user has a valid client_id
            if not client_id:
               return []

        with connection.cursor() as cursor:
            # Execute the stored procedure with the client_id, source_id, and destination_id
            cursor.execute("SELECT * FROM get_combined_rates_data(%s, %s, %s, %s)", [source_id, destination_id, freight_type_id, client_id])
            rows = cursor.fetchall()


            columns = [
                'id', 'unique_uuid', 'company_id', 'company_name', 'rate', 'currency',
                'free_days', 'spot_filed', 'transhipment_add_port', 'effective_date',
                'expiration_date', 'un_number', 'vessel_name', 'cargotype', 'hazardous', 'terms_condition',
                'source_id', 'source_name', 'destination_id', 'destination_name', 'transit_time_id', 'transit_time',
                'freight_type_id', 'freight_type','remarks','shipping_schedule_id','departure_date','arrival_date','port_cut_off_date','si_cut_off_date', 'gate_opening_date','service','voyage','charge','charge_flag','charge_name','pp_cc','note'
            ]

            # print("Columns from DB:", [desc[0] for desc in cursor.description])
            # print("Columns in Django:", columns)

            # columns = config.get("RATE_LIST_QUERYSET" , "").split(",")

            # Convert the result into a list of dictionaries

            data = [dict(zip(columns, row)) for row in rows] #working
            return data


    # def get_queryset(self):
    #     source_id = self.kwargs.get('source')
    #     destination_id = self.kwargs.get('destination')
    #     freight_type_id = self.kwargs.get('freight_type', None)
    #     # [18/feb/25]
    #     user = self.request.user

    #     # Get the client_id of the authenticated user
    #     # client_id = self.request.user.client.client_id if self.request.user.client else None

    #     # if not client_id:
    #     #     raise PermissionDenied({"error": "You are not associated with any client."})

    #     # [18/feb/25]
    #     if user.is_admin or (user.role and user.role.role_name == "Super Admin"):
    #         client_id = None  # Pass NULL in the stored procedure to fetch all data
    #     else: 
    #         client_id = user.client_id

    #     # Check if the user has a valid client_id
    #         if not client_id:
    #            return []

    #     with connection.cursor() as cursor:
    #         # Execute the stored procedure with the client_id, source_id, and destination_id
    #         cursor.execute("SELECT * FROM get_combined_rates_data(%s, %s, %s, %s)", [source_id, destination_id, freight_type_id, client_id])
    #         rows = cursor.fetchall()


    #         columns = [
    #             'id', 'unique_uuid', 'company_id', 'company_name', 'rate', 'currency',
    #             'free_days', 'spot_filed', 'transhipment_add_port', 'effective_date',
    #             'expiration_date', 'un_number', 'vessel_name', 'cargotype', 'hazardous', 'terms_condition',
    #             'source_id', 'source_name', 'destination_id', 'destination_name', 'transit_time_id', 'transit_time',
    #             'freight_type_id', 'freight_type','remarks','shipping_schedule_id','departure_date','arrival_date','port_cut_off_date','si_cut_off_date', 'gate_opening_date','service','voyage','charge','charge_flag','charge_name','pp_cc','note'
    #         ]

    #         # print("Columns from DB:", [desc[0] for desc in cursor.description])
    #         # print("Columns in Django:", columns)

    #         # columns = config.get("RATE_LIST_QUERYSET" , "").split(",")

    #         # Convert the result into a list of dictionaries

    #         data = [dict(zip(columns, row)) for row in rows] #working
    #         return data


# [ MANUALRATE FILTER WITH COMPANY ID ]
class ManualRateFilterWithCompanyIdAPIView(APIView):
    permission_classess =[ IsAuthenticated,
        IsClientUserEditAndRead | IsSystemOrClientAdmin | IsClientUserReadOnly | IsSuperAdmin,]

        #[ 06/MARCH/2025 ]
    def get(self, request, company_id):
        try:
            user = request.user

            # Super Admin can access all manual rates
            if user.is_admin or (user.role and user.role.role_name == "Super Admin"):
                manual_rates = ManualRate.objects.filter(company_id=company_id, soft_delete=False)
            else:
                client_id = user.client_id
                if not client_id:
                    return Response(
                        {"error": "You are not associated with any client."},
                        status=status.HTTP_403_FORBIDDEN
                    )

                manual_rates = ManualRate.objects.filter(
                    company_id=company_id, client_id=client_id, charge='FRTF', soft_delete=False
                ).prefetch_related('shipping_schedules')

                if not manual_rates.exists():
                    return Response(
                        {"error": "FRTF charge code not found for this shipping line."},
                        status=status.HTTP_404_NOT_FOUND
                    )

            serializer = ManualRateSerializer(manual_rates, many=True)
            return Response(serializer.data, status=status.HTTP_200_OK)

        except PermissionDenied as e:
            return Response({"error": str(e)}, status=status.HTTP_403_FORBIDDEN)
        
        except Exception as e:
            return Response(
                {"error": f"An unexpected error occurred: {str(e)}"},
                status=status.HTTP_400_BAD_REQUEST
            )


class ManualRateWithRateWithVersionsAPIView(APIView):
    permission_classes = [
        IsAuthenticated,
        IsClientUserEditAndRead | IsSystemOrClientAdmin | IsClientUserReadOnly | IsSuperAdmin,
    ]

    # [ 11/MAR/2025 ]
    def get(self, request, source_id, destination_id):
        try:
            user = request.user
            logger.debug(f"User: {user}, Role: {getattr(user, 'role', None)}, Current Client: {getattr(user, 'current_client', None)}")

            # Determine the active client for filtering.
            if user.is_admin and user.current_client:
                active_client_id = user.current_client.client_id
            else:
                active_client_id = user.client_id

            if not active_client_id:
                return Response(
                    {"error": "No valid client_id associated with the user."},
                    status=status.HTTP_403_FORBIDDEN
                )

            # For Admin/Super Admin: Fetch data based on the active client
            if user.is_admin or (user.role and user.role.role_name == "Super Admin"):
                manual_rates = ManualRate.objects.filter(
                    source_id=source_id,
                    destination_id=destination_id,
                    client_id=active_client_id,
                    soft_delete=False
                )
            else:
                # For Regular Users: Fetch data only for their assigned client
                manual_rates = ManualRate.objects.filter(
                    source_id=source_id,
                    destination_id=destination_id,
                    client_id=active_client_id,
                    charge='FRTF',
                    soft_delete=False
                )

            # If no matching manual rates, return an appropriate response
            if not manual_rates.exists():
                return Response(
                    {"error": "FRTF charge code not found for this shipping line."},
                    status=status.HTTP_404_NOT_FOUND
                )

            # Serialize results
            serializer = ManualRateSerializer(manual_rates, many=True)
            return Response(serializer.data, status=status.HTTP_200_OK)

        except ManualRate.DoesNotExist:
            return Response(
                {"error": "ManualRate version not found for the given source and destination."},
                status=status.HTTP_404_NOT_FOUND,
            )
        except Exception as e:
            logger.error(f"Unexpected error: {str(e)}", exc_info=True)
            return Response(
                {"error": f"An unexpected error occurred: {str(e)}"},
                status=status.HTTP_400_BAD_REQUEST,
            )

   
class ManualRateListView(APIView):
    permission_classes = [IsAuthenticated, IsClientUserEditAndRead | IsSystemOrClientAdmin | IsClientUserReadOnly | IsSuperAdmin]

    # [11/03/2025]
    def get(self, request):
        try:
            user = request.user

            # Determine active client context for filtering:
            if user.is_admin or (user.role and user.role.role_name == "Super Admin"):
                if user.current_client:
                    # Admin has switched, so filter by the active client's ID.
                    manual_rate = ManualRate.objects.filter(
                        client_id=user.current_client.client_id,
                        soft_delete=False
                    )
                else:
                    # No active switch; admin accesses all manual rates.
                    manual_rate = ManualRate.objects.filter(soft_delete=False)
            else:
                # For regular users, ensure a client is associated.
                client_id = user.current_client.client_id
                if not client_id:
                    return Response(
                        {"error": "You are not associated with any client."},
                        status=status.HTTP_403_FORBIDDEN,
                    )
                manual_rate = ManualRate.objects.filter(
                    client_id=client_id,
                    soft_delete=False
                )

            manual_rate_serializer = ManualRateSerializer(manual_rate, many=True)
            return Response(manual_rate_serializer.data, status=status.HTTP_200_OK)

        except Exception as e:
            return Response(
                {"error": f"An unexpected error occurred: {str(e)}"},
                status=status.HTTP_400_BAD_REQUEST,
            )
    
    # [18/feb/25]
    # def get(self, request):
    #     try:
    #         user = request.user

    #     # Super Admin can access all manual rates
    #         if user.is_admin or (user.role and user.role.role_name == "Super Admin"):
    #             manual_rate = ManualRate.objects.filter(soft_delete=False)
    #         else:
    #             client_id = user.client_id
    #             if not client_id:
    #                 return Response(
    #                 {"error": "You are not associated with any client."},
    #                 status=status.HTTP_403_FORBIDDEN,
    #                 )
    #             manual_rate = ManualRate.objects.filter(client_id=client_id, soft_delete=False)

    #         manual_rate_serializer = ManualRateSerializer(manual_rate, many=True)
    #         return Response(manual_rate_serializer.data, status=status.HTTP_200_OK)

    #     except Exception as e:
    #         return Response(
    #         {"error": f"An unexpected error occurred: {str(e)}"},
    #         status=status.HTTP_400_BAD_REQUEST,
    #     )
    

    # def get(self, request):
    #     # Fetch client_id from the request or user context
    #     user = request.user
    #     client_id = user.client_id if hasattr(user, 'client_id') else None  # Check if user has client_id attribute
    #     if not client_id:
    #         return Response({"detail": "Client ID is required."}, status=status.HTTP_400_BAD_REQUEST)
        
    #     # Filter records by client_id and exclude soft-deleted records
    #     manual_rate = ManualRate.objects.filter(client_id=client_id, soft_delete=False)
    #     manual_rate_serializer = ManualRateSerializer(manual_rate, many=True)
    #     return Response(manual_rate_serializer.data)


    
# UPDATED CODED  [ POST ]

    def post(self, request):
        try:
            with transaction.atomic():
                rate_data = request.data  # Expecting a single dictionary

                # Extract client_id from the request user
                user = request.user
                client_id = getattr(user, 'client_id', None)
                if not client_id:
                    return Response({"detail": "Client ID is required."}, status=status.HTTP_400_BAD_REQUEST)

                # Validate request format
                if not isinstance(rate_data, dict):
                    return Response({"error": "Invalid data format. Expected a single rate data object."}, status=status.HTTP_400_BAD_REQUEST)

                # Required fields
                required_fields = [
                    "company", "source", "destination", "freight_type", "rate", "currency",
                    "effective_date", "expiration_date"
                ]
                for field in required_fields:
                    if field not in rate_data or not rate_data[field]:
                        return Response({"error": f"Missing or empty required field: {field}"}, status=status.HTTP_400_BAD_REQUEST)

                # Extract field values
                company_name = rate_data.get('company')
                source_name = rate_data.get('source')
                destination_name = rate_data.get('destination')
                freight_type_name = rate_data.get('freight_type')
                # transit_time = rate_data.get('transit_time')
                cargotype_name = rate_data.get('cargotype')
                rate = rate_data.get('rate')
                currency = rate_data.get('currency')
                hazardous = rate_data.get('hazardous')
                un_number = rate_data.get('un_number')
                # direct_shipment = rate_data.get('direct_shipment')
                spot_filed = rate_data.get('spot_filed')
                isRateTypeStatus = rate_data.get('isRateTypeStatus')
                vessel_name = rate_data.get('vessel_name')
                voyage = rate_data.get('voyage')
                haz_class = rate_data.get('haz_class')
                packing_group = rate_data.get('packing_group')
                transhipment_add_port = rate_data.get('transhipment_add_port')
                free_days = rate_data.get('free_days')
                free_days_comment = rate_data.get('free_days_comment')
                effective_date = rate_data.get('effective_date')
                expiration_date = rate_data.get('expiration_date')
                remarks = rate_data.get('remarks')
                terms_condition = rate_data.get('terms_condition')
                charge = rate_data.get('charge')
                charge_name = rate_data.get('charge_name')
                charge_flag = rate_data.get('charge_flag')
                pp_cc = rate_data.get('pp_cc')
                note = rate_data.get('note')
                shipping_schedules = rate_data.get('shipping_schedules', [])

                # Retrieve or create related entities
                company_instance, _ = Company.objects.get_or_create(name=company_name,defaults={'client_id': client_id})
                source_instance, _ = Source.objects.get_or_create(name=source_name, defaults={'client_id': client_id})
                destination_instance, _ = Destination.objects.get_or_create(name=destination_name, defaults={'client_id': client_id})
                freight_type_instance, _ = FreightType.objects.get_or_create(type=freight_type_name, defaults={'client_id': client_id})
                # transit_time_instance, _ = TransitTime.objects.get_or_create(time=transit_time, defaults={'client_id': client_id})
                commodity_instance, _ = Comodity.objects.get_or_create(name=cargotype_name, defaults={'client_id': client_id})

                exact_match = ManualRate.objects.filter(
                    client_id=client_id,
                    company=company_instance,
                    source=source_instance,
                    destination=destination_instance,
                    freight_type=freight_type_instance,
                    # transit_time=transit_time_instance,
                    cargotype=commodity_instance,
                    rate=rate,
                    charge=charge,
                    soft_delete=False
                ).exists()

                if exact_match:
                    return Response({"message": "Duplicate record found. No new entry created."}, status=status.HTTP_400_BAD_REQUEST)

                # Generate a unique UUID for the new manual rate
                timestamp = datetime.now().strftime('%Y%m%d%H%M%S%f')
                unique_id = f"{timestamp}{str(uuid.uuid4()).replace('-', '')[:8]}"
                common_uuid = unique_id[:24]  # Use only the first 24 characters of the UUID

                # Create a new record in the ManualRate table
                manual_rate = ManualRate.objects.create(
                    unique_uuid=common_uuid,
                    client_id=client_id,
                    company=company_instance,
                    source=source_instance,
                    destination=destination_instance,
                    freight_type=freight_type_instance,
                    # transit_time=transit_time_instance,
                    cargotype=commodity_instance,
                    rate=rate,
                    currency=currency,
                    hazardous=hazardous,
                    un_number=un_number,
                    # direct_shipment=direct_shipment,
                    spot_filed=spot_filed,
                    isRateTypeStatus=isRateTypeStatus,
                    vessel_name=vessel_name,
                    voyage=voyage,
                    haz_class=haz_class,
                    packing_group=packing_group,
                    transhipment_add_port=transhipment_add_port,
                    free_days=free_days,
                    free_days_comment=free_days_comment,
                    effective_date=effective_date,
                    expiration_date=expiration_date,
                    remarks=remarks,
                    terms_condition=terms_condition,
                    charge=charge,
                    charge_name=charge_name,
                    charge_flag=charge_flag,
                    pp_cc=pp_cc,
                    note=note
                )

                # Create shipping schedules
                
                def parse_date(date_str):
                    try:
                        return datetime.strptime(date_str, '%Y-%m-%d')
                    except ValueError:
                        return None

                for schedule in shipping_schedules:
                    departure_date = schedule.get('departure_date')
                    arrival_date = schedule.get('arrival_date')
                    port_cut_off_date = schedule.get('port_cut_off_date')
                    si_cut_off_date = schedule.get('si_cut_off_date')
                    gate_opening_date = schedule.get('gate_opening_date')
                    service = schedule.get('service')
                    voyage = schedule.get('voyage')

                    # Validate schedule dates
                    if not all([departure_date, arrival_date, port_cut_off_date, si_cut_off_date, gate_opening_date]):
                        raise ValueError("Shipping schedule dates are incomplete.")

                    if departure_date > arrival_date or port_cut_off_date > departure_date or si_cut_off_date > port_cut_off_date or gate_opening_date > si_cut_off_date:
                        raise ValueError("Invalid schedule dates.")

                    # Create ShippingSchedule instance
                    ShippingSchedule.objects.create(
                        manual_rate=manual_rate,
                        departure_date=departure_date,
                        arrival_date=arrival_date,
                        port_cut_off_date=port_cut_off_date,
                        si_cut_off_date=si_cut_off_date,
                        gate_opening_date=gate_opening_date,
                        service=service,
                        voyage=voyage
                    )

                return Response({'message': 'Record created successfully'}, status=status.HTTP_201_CREATED)

        except ValueError as ve:
            return Response({"message": str(ve)}, status=status.HTTP_400_BAD_REQUEST)
        except IntegrityError as e:
            return Response({"detail": f"Database error: {str(e)}"}, status=status.HTTP_400_BAD_REQUEST)
        except Exception as e:
            return Response({"detail": f"An error occurred: {str(e)}"}, status=status.HTTP_400_BAD_REQUEST)


# [ UPDATED ON 29/JAN/25] [ PUT ]
    def put(self, request ,unique_uuid):
        try:
            with transaction.atomic():
                rate_data = request.data  # Expecting a dictionary with updated data

                # Extract client_id from request user
                user = request.user
                client_id = getattr(user, 'client_id', None)
                if not client_id:
                    return Response({"detail": "Client ID is required."}, status=status.HTTP_400_BAD_REQUEST)

                # Validate request format
                if not isinstance(rate_data, dict):
                    return Response({"error": "Invalid data format. Expected a single rate data object."}, status=status.HTTP_400_BAD_REQUEST)

                # Check if unique_uuid is provided
                # unique_uuid = rate_data.get('unique_uuid')
                unique_uuid = unique_uuid
                if not unique_uuid:
                    return Response({"error": "Unique UUID is required for updating a record."}, status=status.HTTP_400_BAD_REQUEST)

                # Retrieve the existing ManualRate record
                try:
                    manual_rate = ManualRate.objects.get(unique_uuid=unique_uuid, client_id=client_id)
                except ManualRate.DoesNotExist:
                    return Response({"error": "Record not found."}, status=status.HTTP_404_NOT_FOUND)

                # Required fields for validation
                required_fields = [
                    "company", "source", "destination", "freight_type", "rate", "currency",
                    "effective_date", "expiration_date"
                ]
                for field in required_fields:
                    if field not in rate_data or not rate_data[field]:
                        return Response({"error": f"Missing or empty required field: {field}"}, status=status.HTTP_400_BAD_REQUEST)

                # Extract updated field values
                company_name = rate_data.get('company')
                source_name = rate_data.get('source')
                destination_name = rate_data.get('destination')
                freight_type_name = rate_data.get('freight_type')
                # transit_time = rate_data.get('transit_time')
                cargotype_name = rate_data.get('cargotype')
                rate = rate_data.get('rate')
                currency = rate_data.get('currency')
                hazardous = rate_data.get('hazardous')
                un_number = rate_data.get('un_number')
                direct_shipment = rate_data.get('direct_shipment')
                spot_filed = rate_data.get('spot_filed')
                isRateTypeStatus = rate_data.get('isRateTypeStatus')
                vessel_name = rate_data.get('vessel_name')
                voyage = rate_data.get('voyage')
                haz_class = rate_data.get('haz_class')
                packing_group = rate_data.get('packing_group')
                transhipment_add_port = rate_data.get('transhipment_add_port')
                free_days = rate_data.get('free_days')
                free_days_comment = rate_data.get('free_days_comment')
                effective_date = rate_data.get('effective_date')
                expiration_date = rate_data.get('expiration_date')
                remarks = rate_data.get('remarks')
                terms_condition = rate_data.get('terms_condition')
                charge = rate_data.get('charge')
                charge_name = rate_data.get('charge_name')
                charge_flag = rate_data.get('charge_flag')
                pp_cc = rate_data.get('pp_cc')
                note = rate_data.get('note')
                shipping_schedules = rate_data.get('shipping_schedules', [])

                # Retrieve or create related entities
                company_instance, _ = Company.objects.get_or_create(name=company_name, defaults={'client_id': client_id})
                source_instance, _ = Source.objects.get_or_create(name=source_name, defaults={'client_id': client_id})
                destination_instance, _ = Destination.objects.get_or_create(name=destination_name, defaults={'client_id': client_id})
                freight_type_instance, _ = FreightType.objects.get_or_create(type=freight_type_name, defaults={'client_id': client_id})
                # transit_time_instance, _ = TransitTime.objects.get_or_create(time=transit_time, defaults={'client_id': client_id})
                commodity_instance, _ = Comodity.objects.get_or_create(name=cargotype_name, defaults={'client_id': client_id})

                # Check if an exact updated record already exists
                exact_match = ManualRate.objects.filter(
                    client_id=client_id,
                    company=company_instance,
                    source=source_instance,
                    destination=destination_instance,
                    freight_type=freight_type_instance,
                    # transit_time=transit_time_instance,
                    cargotype=commodity_instance,
                    rate=rate,
                    charge=charge,
                    soft_delete=False

                ).exclude(unique_uuid=unique_uuid).exists()

                if exact_match:
                    return Response({"message": "Duplicate record found. No changes were made."}, status=status.HTTP_400_BAD_REQUEST)

                # Update fields only if values have changed
                manual_rate.company = company_instance
                manual_rate.source = source_instance
                manual_rate.destination = destination_instance
                manual_rate.freight_type = freight_type_instance
                # manual_rate.transit_time = transit_time_instance
                manual_rate.cargotype = commodity_instance.name
                manual_rate.rate = rate
                manual_rate.currency = currency
                manual_rate.hazardous = hazardous
                manual_rate.un_number = un_number
                # manual_rate.direct_shipment = direct_shipment
                manual_rate.spot_filed = spot_filed
                manual_rate.isRateTypeStatus = isRateTypeStatus
                manual_rate.vessel_name = vessel_name
                manual_rate.voyage = voyage
                manual_rate.haz_class = haz_class
                manual_rate.packing_group = packing_group
                manual_rate.transhipment_add_port = transhipment_add_port
                manual_rate.free_days = free_days
                manual_rate.free_days_comment = free_days_comment
                manual_rate.effective_date = effective_date
                manual_rate.expiration_date = expiration_date
                manual_rate.remarks = remarks
                manual_rate.terms_condition = terms_condition
                manual_rate.charge=charge
                manual_rate.charge_name=charge_name
                manual_rate.charge_flag=charge_flag
                manual_rate.pp_cc=pp_cc
                manual_rate.note=note
                manual_rate.save()

                # Update shipping schedules
                ShippingSchedule.objects.filter(manual_rate=manual_rate).delete()  # Remove old schedules
                for schedule in shipping_schedules:
                    departure_date = schedule.get('departure_date')
                    arrival_date = schedule.get('arrival_date')
                    port_cut_off_date = schedule.get('port_cut_off_date')
                    si_cut_off_date = schedule.get('si_cut_off_date')
                    gate_opening_date = schedule.get('gate_opening_date')
                    service = schedule.get('service')
                    voyage = schedule.get('voyage')

                    # Validate schedule dates
                    if not all([departure_date, arrival_date, port_cut_off_date, si_cut_off_date, gate_opening_date]):
                        raise ValueError("Shipping schedule dates are incomplete.")

                    if departure_date > arrival_date or port_cut_off_date > departure_date or si_cut_off_date > port_cut_off_date or gate_opening_date > si_cut_off_date:
                        raise ValueError("Invalid schedule dates.")

                    # Create new ShippingSchedule entries
                    ShippingSchedule.objects.create(
                        manual_rate=manual_rate,
                        departure_date=departure_date,
                        arrival_date=arrival_date,
                        port_cut_off_date=port_cut_off_date,
                        si_cut_off_date=si_cut_off_date,
                        gate_opening_date=gate_opening_date,
                        service=service,
                        voyage=voyage
                    )

                return Response({'message': 'Record updated successfully'}, status=status.HTTP_200_OK)

        except ValueError as ve:
            return Response({"message": str(ve)}, status=status.HTTP_400_BAD_REQUEST)
        except IntegrityError as e:
            return Response({"detail": f"Database error: {str(e)}"}, status=status.HTTP_400_BAD_REQUEST)
        except Exception as e:
            return Response({"detail": f"An error occurred: {str(e)}"}, status=status.HTTP_400_BAD_REQUEST) 



    def delete(self, request, unique_uuid):
        try:
            user = request.user
            client_id = user.client_id if hasattr(user, 'client_id') else None
            if not client_id:
                return Response({"detail": "Client ID is required."}, status=status.HTTP_400_BAD_REQUEST)

            # Retrieve and soft-delete the ManualRate record
            manual_rate_instance = ManualRate.objects.get(unique_uuid=unique_uuid, client_id=client_id, soft_delete=False)
            manual_rate_instance.soft_delete = True
            manual_rate_instance.save()

            return Response({'message': 'ManualRate soft-deleted successfully'}, status=status.HTTP_204_NO_CONTENT)

        except ManualRate.DoesNotExist:
            return Response({"detail": "ManualRate not found."}, status=status.HTTP_404_NOT_FOUND)
        
        except Exception as e:
            return Response({'detail': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
      

class UpdatingRateFrozenInfoListView(APIView):
    permission_classes = [IsAuthenticated, IsClientUserEditAndRead | IsSystemOrClientAdmin | IsClientUserReadOnly]

    def put(self, request, unique_uuid):
        try:
            requestData = request.data
            isRateUsed = requestData.get('isRateUsed', None)
            client_id = request.user.client_id  # Assuming client_id is associated with the authenticated user.

            if isRateUsed is None:
                return Response({"error": "The 'isRateUsed' field is required."}, status=status.HTTP_400_BAD_REQUEST)

            # Fetch the ManualRate object with the client_id filter
            try:
                manual_rate = ManualRate.objects.get(unique_uuid=unique_uuid, client_id=client_id)
            except ManualRate.DoesNotExist:
                return Response({"error": "ManualRate not found or access denied."}, status=status.HTTP_404_NOT_FOUND)

            # Update the isRateUsed field
            manual_rate.isRateUsed = isRateUsed
            manual_rate.save()

            return Response({"message": "The 'isRateUsed' status has been updated successfully."}, status=status.HTTP_200_OK)

        except PermissionDenied:
            return Response({"error": "Permission denied. You are not authorized to update this record."}, status=status.HTTP_403_FORBIDDEN)
        except ObjectDoesNotExist:
            return Response({"error": "The requested resource does not exist."}, status=status.HTTP_404_NOT_FOUND)
        except Exception as e:
            return Response({"error": f"An unexpected error occurred: {str(e)}"}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
                

class CustomerInfoListView(APIView):
    permission_classes = [IsAuthenticated, IsSystemOrClientAdmin | IsClientUserEditAndRead | IsSuperAdmin,]

    # FOR GET
    # [ 18/feb/25] 
    def get(self, request):
        try:
            user = request.user
            
            # Super Admin can access all sources without filtering by client_id
            if user.is_admin or (user.role and user.role.role_name == "Super Admin"):
                # Super Admins are not restricted by client_id
                freightType = CustomerInfo.objects.all()  # Get all customer info, no filter needed
            else:
                # Non-admin users are filtered by their client_id
                freightType = CustomerInfo.objects.filter(client_id=user.client_id)

            # Serialize the data and return it
            serializer = CustomerInfoSerializer(freightType, many=True)
            return Response(serializer.data, status=status.HTTP_200_OK)

        except Exception as e:
            return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


    # def get(self, request):
    #     try:
    #         client_id = request.user.client_id  # Assuming the user is associated with a client_id
    #         customer_info = CustomerInfo.objects.filter(client_id=client_id)

    #         if not customer_info.exists():
    #             return Response({"message": "No customer information found."}, status=status.HTTP_404_NOT_FOUND)

    #         customer_info_serializer = CustomerInfoSerializer(customer_info, many=True)
    #         return Response(customer_info_serializer.data, status=status.HTTP_200_OK)

    #     except PermissionDenied:
    #         return Response({"error": "You do not have permission to access this resource."}, status=status.HTTP_403_FORBIDDEN)
    #     except Exception as e:
    #         return Response({"error": f"An unexpected error occurred: {str(e)}"}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    # FOR POST
    def post(self, request):
        try:
            # requestData = request.data

            requestData = request.data
            client_id = request.user.client_id
                #new added at 24 jan
                # logger = logging.getLogger(_name_)
            logger.info(f"Logged-in user: {request.user}, client_id: {client_id}")
            if not client_id:
                return Response({"detail": "Client ID is required."}, status=status.HTTP_400_BAD_REQUEST)
            

            # Validate required fields
            required_fields = ['company_name', 'cust_name', 'cust_email', 'sales_represent', 'phone', 'percentage', 'terms_condition']
            for field in required_fields:
                if not requestData.get(field):
                    return Response({"error": f"{field} is required."}, status=status.HTTP_400_BAD_REQUEST)

            # Create the CustomerInfo instance
            customer_serializer = CustomerInfo.objects.create(
                client_id=request.user.client_id,  # Associate with the client's ID
                company_name=requestData.get('company_name'),
                cust_name=requestData.get('cust_name'),
                cust_email=requestData.get('cust_email'),
                sales_represent=requestData.get('sales_represent'),
                phone=requestData.get('phone'),
                percentage=requestData.get('percentage'),
                terms_condition=requestData.get('terms_condition'),
            )
            severResponse = {
                'id':customer_serializer.id,
                'client_id':client_id,
                'company_name':customer_serializer.company_name,
                'cust_name':customer_serializer.cust_name,
                'cust_email':customer_serializer.cust_email,
                'sales_represent':customer_serializer.sales_represent,
                'phone':customer_serializer.phone,
                'percentage':customer_serializer.percentage,
                'terms_condition':customer_serializer.terms_condition,
            }

            return Response({'message': 'Customer created successfully' , 'data': severResponse}, status=status.HTTP_201_CREATED)

        except PermissionDenied:
            return Response({"error": "You do not have permission to perform this action."}, status=status.HTTP_403_FORBIDDEN)
        except Exception as e:
            return Response({'error': f"An unexpected error occurred: {str(e)}"}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    # FOR UPDATE BY ID
    def put(self, request, id):
        try:
            requestData = request.data
            client_id = request.user.client_id  # Assuming the user is associated with a client_id

            try:
                customer_instance = CustomerInfo.objects.get(id=id, client_id=client_id)
            except CustomerInfo.DoesNotExist:
                return Response({"error": "Customer not found or unauthorized access."}, status=status.HTTP_404_NOT_FOUND)

            # Update fields if provided
            if 'percentage' in requestData:
                customer_instance.percentage = requestData['percentage']
            
            # You can add more fields to update here as required

            customer_instance.save()

            return Response({"message": "Customer information updated successfully."}, status=status.HTTP_200_OK)

        except PermissionDenied:
            return Response({"error": "You do not have permission to update this resource."}, status=status.HTTP_403_FORBIDDEN)
        except Exception as e:
            return Response({'error': f"An unexpected error occurred: {str(e)}"}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
  

class CustomerInfoDetailsListView(APIView):
    permission_classes = [IsAuthenticated, IsSystemOrClientAdmin | IsClientUserEditAndRead | IsSuperAdmin,]
    
    # FOR GET BY ID
    #[11/mar/25]
    def get(self, request, id):
        try:
            user = request.user

            if user.is_admin or (user.role and user.role.role_name == "Super Admin"):
                if user.current_client:
                    # If admin has switched, restrict to the active client.
                    try:
                        customer_info = CustomerInfo.objects.get(
                            id=id, client_id=user.current_client.client_id
                        )
                    except CustomerInfo.DoesNotExist:
                        return Response(
                            {"error": f"CustomerInfo with id {id} does not exist or unauthorized access."},
                            status=status.HTTP_404_NOT_FOUND,
                        )
                else:
                    # Admin without an active switch can access all records.
                    customer_info = CustomerInfo.objects.get(id=id)
            else:
                client_id = user.client_id
                if not client_id:
                    return Response(
                        {"error": "You are not associated with any client."},
                        status=status.HTTP_403_FORBIDDEN,
                    )
                try:
                    customer_info = CustomerInfo.objects.get(
                        id=id, client_id=client_id
                    )
                except CustomerInfo.DoesNotExist:
                    return Response(
                        {"error": f"CustomerInfo with id {id} does not exist or unauthorized access."},
                        status=status.HTTP_404_NOT_FOUND,
                    )

            serializer = CustomerInfoSerializer(customer_info)
            return Response(serializer.data, status=status.HTTP_200_OK)

        except Exception as e:
            return Response(
                {"error": f"An unexpected error occurred: {str(e)}"},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR,
            )

    #[18/feb/25]
    # def get(self, request, id):
    #     try:
    #         user = request.user

    #         # Super Admin can access all customer info
    #         if user.is_admin or (user.role and user.role.role_name == "Super Admin"):
    #             customer_info = CustomerInfo.objects.get(id=id)
    #         else:
    #             client_id = user.client_id
    #             if not client_id:
    #                 return Response(
    #                     {"error": "You are not associated with any client."},
    #                     status=status.HTTP_403_FORBIDDEN,
    #                 )

    #             # Fetch customer info for the given id and client_id
    #             try:
    #                 customer_info = CustomerInfo.objects.get(id=id, client_id=client_id)
    #             except CustomerInfo.DoesNotExist:
    #                 return Response(
    #                     {"error": f"CustomerInfo with id {id} does not exist or unauthorized access."},
    #                     status=status.HTTP_404_NOT_FOUND,
    #                 )

    #         # Move this outside the if-else block
    #         customer_info_serializer = CustomerInfoSerializer(customer_info)
    #         return Response(customer_info_serializer.data, status=status.HTTP_200_OK)

    #     except CustomerInfo.DoesNotExist:
    #         return Response(
    #             {"error": f"CustomerInfo with id {id} not found."},
    #             status=status.HTTP_404_NOT_FOUND,
    #         )
    #     except Exception as e:
    #         return Response(
    #             {"error": f"An unexpected error occurred: {str(e)}"},
    #             status=status.HTTP_500_INTERNAL_SERVER_ERROR,
    #         )



class CommodityList(generics.ListCreateAPIView):
    permission_classes = [IsAuthenticated, IsClientUserEditAndRead | IsSystemOrClientAdmin | IsClientUserReadOnly | IsUser | IsSuperAdmin,]
    serializer_class = CommoditySerializer

    #[11/mar/25]
    def get_queryset(self):
        user = self.request.user

        if user.is_admin or (user.role and user.role.role_name == "Super Admin"):
            # If admin has switched, use the active client filter.
            if user.current_client:
                return Comodity.objects.filter(client_id=user.current_client.client_id)
            # Otherwise, return all commodities.
            return Comodity.objects.all()

        # For regular users, ensure client_id exists and filter accordingly.
        if hasattr(user, 'client_id') and user.client_id:
            return Comodity.objects.filter(client_id=user.client_id)

        return Comodity.objects.none()


  
    #[18/feb/25] 
    # def get_queryset(self):
    
    #     user = self.request.user

    # # Super Admin can access all commodities
    #     if user.is_admin or (user.role and user.role.role_name == "Super Admin"):
    #         return Comodity.objects.all()

    #     if hasattr(user, 'client_id') and user.client_id:
    #         return Comodity.objects.filter(client_id=user.client_id)

    #     return Comodity.objects.none()
    
  

    def perform_create(self, serializer):
        """
        Automatically associate the created commodity with the user's client, if applicable.
        """
        user = self.request.user
        try:
            if hasattr(user, 'client_id'):
                serializer.save(client_id=user.client_id)  # Save the commodity with the associated client_id
            else:
                serializer.save()  # Admin/system users can save without a client_id
        except Exception as e:
            return Response({"error": f"An error occurred while creating the commodity: {str(e)}"}, 
                            status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    def post(self, request, *args, **kwargs):
        """
        Handles commodity creation with proper error handling and response messages.
        """
        try:
            return super().post(request, *args, **kwargs)
        except PermissionDenied:
            return Response({"error": "You do not have permission to create this commodity."}, 
                            status=status.HTTP_403_FORBIDDEN)
        except Exception as e:
            return Response({"error": f"An unexpected error occurred while creating the commodity: {str(e)}"}, 
                            status=status.HTTP_500_INTERNAL_SERVER_ERROR)


class IncoTermList(generics.ListCreateAPIView):
    permission_classes = [IsAuthenticated, IsClientUserEditAndRead | IsSystemOrClientAdmin | IsClientUserReadOnly | IsUser | IsSuperAdmin,]
    serializer_class = IncoTermSerializer

    # [11/mar/25]
    def get_queryset(self):
        user = self.request.user

        if user.is_admin or (user.role and user.role.role_name == "Super Admin"):
            if user.current_client:
                return IncoTerm.objects.filter(client_id=user.current_client.client_id)
            return IncoTerm.objects.all()

        if hasattr(user, 'client_id') and user.client_id:
            return IncoTerm.objects.filter(client_id=user.client_id)

        return IncoTerm.objects.none()

    # [18/feb/25]
    # def get_queryset(self):
    #     user = self.request.user

    # # Super Admin can access all IncoTerms
    #     if user.is_admin or (user.role and user.role.role_name == "Super Admin"):
    #         return IncoTerm.objects.all()

    #     if hasattr(user, 'client_id') and user.client_id:
    #         return IncoTerm.objects.filter(client_id=user.client_id)

    #     return IncoTerm.objects.none() 



    def perform_create(self, serializer):
        """
        Automatically associate the created inco term with the user's client, if applicable.
        """
        user = self.request.user
        try:
            if hasattr(user, 'client_id'):
                serializer.save(client_id=user.client_id)  # Save the inco term with the associated client_id
            else:
                serializer.save()  # Admin/system users can save without a client_id
        except Exception as e:
            return Response({"error": f"An error occurred while creating the inco term: {str(e)}"}, 
                            status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    def post(self, request, *args, **kwargs):
        """
        Handles inco term creation with proper error handling and response messages.
        """
        try:
            return super().post(request, *args, **kwargs)
        except PermissionDenied:
            return Response({"error": "You do not have permission to create this inco term."}, 
                            status=status.HTTP_403_FORBIDDEN)
        except Exception as e:
            return Response({"error": f"An unexpected error occurred while creating the inco term: {str(e)}"}, 
                            status=status.HTTP_500_INTERNAL_SERVER_ERROR)


# ACTIVITY LOG FUNCTION

# 30/Dec/2024
class ActivityLogView(APIView):
    permission_classes = [IsAuthenticated,IsClientUserEditAndRead|IsSystemOrClientAdmin|IsClientUserReadOnly|IsUser | IsSuperAdmin]  # Ensure only authenticated users access this view

    def get(self, request):
        try:
            # Fetch recent_only filter from query parameters
            recent_only = request.query_params.get('recent', 'false').lower() == 'true'

            # Filter logs for the logged-in user only
            query = ActivityLog.objects.filter(user=request.user)

            # If recent_only is enabled, fetch logs for the past 7 days
            if recent_only:
                seven_days_ago = datetime.now() - timedelta(days=7)
                query = query.filter(created_at__gte=seven_days_ago)

            # Order logs by creation date (most recent first)
            activitityLogList = query.order_by('-created_at')

            # Serialize and return the data
            activitityLogListSerializer = ActivityLogSerializer(activitityLogList, many=True)
            return Response(activitityLogListSerializer.data)

        except Exception as err:
            return Response({"error": "Something went wrong", "details": str(err)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    def post(self, request):
        try:
            with transaction.atomic():
                requestData = request.data
                
                # Create an activity log linked to the logged-in user
                ActivityLog.objects.create(
                    user=request.user,
                    action_type=requestData.get('action_type'),
                    action_status=requestData.get('action_status'),
                    description=requestData.get('description'),
                    source_id=requestData.get('source_id'),
                    destination_id=requestData.get('destination_id'),
                )
            return Response({'message': 'Log created successfully'}, status=status.HTTP_201_CREATED)

        except Exception as err:
            return Response({"error": "Something went wrong", "details": str(err)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

# CLIENT INFO 
class ClientinfoViewSet(APIView):
    
    #[GET]
    def get(self, request):
        try:
            client_info = Clientinfo.objects.all()
            client_info_serializer = ClientinfoSerializer(client_info, many=True)
            return Response(client_info_serializer.data)

        except Exception as err:
            return Response({"error": "Something went wrong", "details": str(err)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    #[POST]
    def post(self, request):
        try:
            # # Extract request data
            requestData = request.data

            # # Generate unique ID
            timestamp = datetime.now().strftime('%Y%m%d%H%M%S')
            company_name = requestData.get('company_name', '').strip()  # Ensure no leading/trailing spaces
            if not company_name:
                return Response({"error": "Company name is required"}, status=status.HTTP_400_BAD_REQUEST)

            # Split company name into parts
            company_name_parts = company_name.split()
            if len(company_name_parts) == 0:
                return Response({"error": "Invalid company name format"}, status=status.HTTP_400_BAD_REQUEST)

            # Capitalize the first word and format the rest
            company_name_parts[0] = company_name_parts[0].capitalize()
            formatted_company_name = "_".join([company_name_parts[0]] + [word.lower() for word in company_name_parts[1:]])
            unique_id = f"{formatted_company_name}_{timestamp}"

            # Extract other fields from request
            clientName = requestData.get('client_name')
            companyName = requestData.get('company_name')
            email = requestData.get('email')
            address = requestData.get('address')
            phoneNo = requestData.get('phone_no')
            invoicingCurrency = requestData.get('invoicing_currency')
            reportingCurrency = requestData.get('reporting_currency')
            region = requestData.get('region')

                # Validation for required fields
                # if not all([client_name, email, address, phone_no, invoicing_currency, reporting_currency, region]):
                #     return Response({"error": "All fields are required"}, status=status.HTTP_400_BAD_REQUEST)

                # Create an activity log linked to the logged-in user
                # Uncomment this block to save the data
            with transaction.atomic():    
                Clientinfo.objects.create(
                    client_id=unique_id,
                    client_name=clientName,
                    company_name=companyName,
                    email=email,
                    address=address,
                    phone_no=phoneNo,
                    invoicing_currency=invoicingCurrency,
                    reporting_currency=reportingCurrency,
                    region=region,
                    created_at=now()
                )

            return Response({'message': unique_id}, status=status.HTTP_201_CREATED)

        except Exception as err:
            return Response({"error": "Something went wrong", "details": str(err)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        

#[ RECENT QUOTATION GENERATED ]
class RecentQuotationsView(generics.ListAPIView):
    serializer_class = QuotationSerializer
    permission_classes = [permissions.IsAuthenticated]

    def get_queryset(self):
        user_client = self.request.user.client_id  # Assuming Client_id is linked to User
        rates = Rate.objects.filter(client=user_client, soft_delete=False).order_by('-effective_date')[:10]
        versioned_rates = VersionedRate.objects.filter(client=user_client, soft_delete=False).order_by('-effective_date')[:10]
        manual_rates = ManualRate.objects.filter(client=user_client, soft_delete=False).order_by('-effective_date')[:10]
        
        # Combine all quotations
        quotations = list(rates) + list(versioned_rates) + list(manual_rates)
        return sorted(quotations, key=lambda x: x.effective_date, reverse=True)[:10]
